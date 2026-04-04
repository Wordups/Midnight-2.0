"""
================================================================================
HPS Policy Migration Builder  —  Strict Layout Engine v2
================================================================================
What was rebuilt:
  - Section content rows no longer use prevent_row_break_across_pages
    (this was the main cause of layout breaks on long content)
  - Banner row uses exact 720-twip height; logo fixed at 2.2" width
  - All column widths locked in fixed layout; never percentage-driven
  - Applicable To left cell: explicit paragraph + valign alignment
  - Revision history: handles tuple, list, and dict entries; never skipped
  - Footer: pinned at 7.0" centered; borders cleared on both rows
  - Logo path: graceful fallback — missing file renders text, never crashes
  - rembg: not imported here; logo processing lives in app.py
================================================================================
"""

from __future__ import annotations

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Page geometry (twips) ──────────────────────────────────────────────────────
_PAGE_W_IN   = 8.5
_MARGIN_IN   = 0.75
CONTENT_IN   = _PAGE_W_IN - 2 * _MARGIN_IN          # 7.0"
CONTENT_W    = int(CONTENT_IN * 1440)                 # 10 080 twips

# ── Brand colors ───────────────────────────────────────────────────────────────
GRAY_BANNER  = "BFBFBF"
GRAY_LABEL   = "D9D9D9"
GRAY_SUBHDR  = "BFBFBF"
GRAY_SECTION = "D9D9D9"
WHITE        = "FFFFFF"
BLACK        = "000000"
WIPRO_RED    = "C00000"
WIPRO_NAVY   = "17375E"
FOOTER_GRAY  = "595959"

# ── Default logo (override via logo_path argument) ─────────────────────────────
DEFAULT_LOGO_PATH = ""


# ══════════════════════════════════════════════════════════════════════════════
# LOW-LEVEL XML HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _rgb(hex_str: str):
    h = hex_str.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def _remove(parent, tag):
    for old in parent.findall(qn(tag)):
        parent.remove(old)


def styled_run(para, text, bold=False, italic=False, underline=False,
               color_hex=BLACK, size_pt=9.5, font="Arial"):
    run = para.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.underline = underline
    run.font.name = font
    run.font.size = Pt(size_pt)
    r, g, b = _rgb(color_hex)
    run.font.color.rgb = RGBColor(r, g, b)
    return run


def _para_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    _remove(pPr, "w:spacing")
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:before"), str(before))
    spc.set(qn("w:after"),  str(after))
    if line is not None:
        spc.set(qn("w:line"),     str(line))
        spc.set(qn("w:lineRule"), "auto")
    pPr.append(spc)


def _keep_with_next(para, on=True):
    pPr = para._p.get_or_add_pPr()
    _remove(pPr, "w:keepNext")
    if on:
        pPr.append(OxmlElement("w:keepNext"))


def _keep_lines(para, on=True):
    pPr = para._p.get_or_add_pPr()
    _remove(pPr, "w:keepLines")
    if on:
        pPr.append(OxmlElement("w:keepLines"))


def _cell_shade(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    _remove(tcPr, "w:shd")
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex.upper())
    tcPr.append(shd)


def _cell_borders(cell, color="000000", size=4):
    tcPr = cell._tc.get_or_add_tcPr()
    _remove(tcPr, "w:tcBorders")
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def _cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    _remove(tcPr, "w:tcMar")
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def _cell_valign(cell, align=WD_ALIGN_VERTICAL.CENTER):
    tcPr = cell._tc.get_or_add_tcPr()
    _remove(tcPr, "w:vAlign")
    va = OxmlElement("w:vAlign")
    va.set(qn("w:val"), "center" if align == WD_ALIGN_VERTICAL.CENTER else "top")
    tcPr.append(va)


def _col_width(cell, twips):
    tcPr = cell._tc.get_or_add_tcPr()
    _remove(tcPr, "w:tcW")
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _row_height(row, twips, exact=False):
    trPr = row._tr.get_or_add_trPr()
    _remove(trPr, "w:trHeight")
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"),   str(int(twips)))
    trH.set(qn("w:hRule"), "exact" if exact else "atLeast")
    trPr.append(trH)


def _no_row_break(row):
    trPr = row._tr.get_or_add_trPr()
    _remove(trPr, "w:cantSplit")
    trPr.append(OxmlElement("w:cantSplit"))


def _style_cell(cell, shade=WHITE, border_color="000000", border_size=4, margins=True):
    _cell_shade(cell, shade)
    _cell_borders(cell, color=border_color, size=border_size)
    if margins:
        _cell_margins(cell)
    return cell


# ══════════════════════════════════════════════════════════════════════════════
# TABLE FACTORY  —  fixed layout, locked column widths
# ══════════════════════════════════════════════════════════════════════════════

def _new_table(doc, rows, cols, col_widths, total_width):
    """Create a fixed-layout table.  col_widths must sum to total_width."""
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"

    tblPr = tbl._tbl.tblPr
    _remove(tblPr, "w:tblW")
    _remove(tblPr, "w:tblLayout")

    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    str(total_width))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)

    tblGrid = tbl._tbl.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        tbl._tbl.insert(0, tblGrid)
    for old in tblGrid.findall(qn("w:gridCol")):
        tblGrid.remove(old)
    for w in col_widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tblGrid.append(gc)

    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            _col_width(cell, col_widths[min(i, len(col_widths) - 1)])

    return tbl


# ══════════════════════════════════════════════════════════════════════════════
# PARAGRAPH BUILDERS
# ══════════════════════════════════════════════════════════════════════════════

def _label_para(cell, text, size_pt=9.0):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _para_spacing(para, 30, 30)
    for i, line in enumerate(text.split("\n")):
        if i:
            para.add_run("\n")
        styled_run(para, line, bold=True, size_pt=size_pt)
    return para


def _value_para(cell, text, size_pt=9.5, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    _para_spacing(para, 30, 30)
    styled_run(para, text or "", size_pt=size_pt, bold=bold)
    return para


def _center_bold_para(cell, text, size_pt=9.5, color_hex=BLACK):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(para, 40, 40)
    styled_run(para, text, bold=True, size_pt=size_pt, color_hex=color_hex)
    return para


def _section_hdr_para(cell, text, size_pt=10.0):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _para_spacing(para, 30, 30)
    _keep_with_next(para, True)
    styled_run(para, text, bold=True, size_pt=size_pt)
    return para


def _content_para(cell, text, before=40, after=30, size_pt=10.0,
                  bold_prefix=None, italic_prefix=None):
    para = cell.add_paragraph()
    _para_spacing(para, before, after)
    _keep_lines(para, True)
    if italic_prefix:
        styled_run(para, italic_prefix[0], bold=True, italic=True, size_pt=size_pt)
        styled_run(para, italic_prefix[1], size_pt=size_pt)
    elif bold_prefix:
        styled_run(para, bold_prefix[0], bold=True, size_pt=size_pt)
        styled_run(para, bold_prefix[1], size_pt=size_pt)
    else:
        styled_run(para, text, size_pt=size_pt)
    return para


def _heading_para(cell, text, size_pt=10.0):
    para = cell.add_paragraph()
    _para_spacing(para, 60, 30)
    _keep_with_next(para, True)
    styled_run(para, text, bold=True, underline=True, size_pt=size_pt)
    return para


def _empty_para(cell, before=0, after=0):
    para = cell.add_paragraph()
    _para_spacing(para, before, after)
    return para


def _bullet_para(cell, text, is_sub=False):
    para = cell.add_paragraph()
    _para_spacing(para, 30, 30)
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"),    str(900 if is_sub else 540))
    ind.set(qn("w:hanging"), "260")
    pPr.append(ind)
    char = "\u25E6" if is_sub else "\u2022"
    styled_run(para, f"{char}  {text}", size_pt=10)
    return para


def _semi_breaks(cell, text, size_pt=10.0):
    segs = [s.strip() for s in text.split(";") if s.strip()]
    for i, seg in enumerate(segs):
        suffix = ";" if i < len(segs) - 1 else ""
        para = cell.add_paragraph()
        _para_spacing(para, 30, 30)
        styled_run(para, seg + suffix, size_pt=size_pt)


# ══════════════════════════════════════════════════════════════════════════════
# REVISION HISTORY NORMALIZER
# ══════════════════════════════════════════════════════════════════════════════

def _norm_revision(entry):
    """Accept tuple, list, or dict.  Always returns a 4-tuple of strings."""
    if isinstance(entry, dict):
        return (
            str(entry.get("date",        "")),
            str(entry.get("version",     "")),
            str(entry.get("updated_by",  "")),
            str(entry.get("description", "")),
        )
    if isinstance(entry, (list, tuple)):
        padded = list(entry) + ["", "", "", ""]
        return tuple(str(x) for x in padded[:4])
    # Last-resort: stringify whatever arrived
    return (str(entry), "", "", "")


# ══════════════════════════════════════════════════════════════════════════════
# MAIN BUILDER
# ══════════════════════════════════════════════════════════════════════════════

def build_policy_document(data: dict, output_path: str, logo_path: str | None = None):
    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────────
    sec = doc.sections[0]
    sec.page_width      = Inches(8.5)
    sec.page_height     = Inches(11)
    sec.left_margin     = Inches(_MARGIN_IN)
    sec.right_margin    = Inches(_MARGIN_IN)
    sec.top_margin      = Inches(0.75)
    sec.bottom_margin   = Inches(0.9)
    sec.footer_distance = Inches(0.3)

    doc.styles["Normal"].paragraph_format.space_before = Pt(0)
    doc.styles["Normal"].paragraph_format.space_after  = Pt(0)

    W = CONTENT_W   # 10 080 twips  (alias for readability below)

    def _gap(before=36, after=0):
        p = doc.add_paragraph()
        _para_spacing(p, before, after)

    effective_logo = logo_path or DEFAULT_LOGO_PATH

    # ──────────────────────────────────────────────────────────────────────────
    # TOP TABLE  —  banner + all metadata rows
    # ──────────────────────────────────────────────────────────────────────────
    # Column widths (must sum to W = 10 080)
    CL  = int(W * 0.18)   # 1 814  — label left
    CM  = int(W * 0.32)   # 3 225  — value left
    CRL = int(W * 0.18)   # 1 814  — label right
    CRV = W - CL - CM - CRL  # remainder — value right

    top = _new_table(doc, 12, 4, [CL, CM, CRL, CRV], W)

    # — Row 0: Banner ──────────────────────────────────────────────────────────
    banner = top.rows[0].cells[0]
    banner = banner.merge(top.rows[0].cells[3])
    _style_cell(banner, GRAY_BANNER)
    _cell_margins(banner, top=20, bottom=20, left=20, right=20)
    _cell_valign(banner, WD_ALIGN_VERTICAL.CENTER)
    _row_height(top.rows[0], 720, exact=True)   # 0.5" — never expands
    _no_row_break(top.rows[0])

    banner.text = ""
    bp = banner.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(bp, 0, 0)

    _logo_ok = effective_logo and os.path.exists(str(effective_logo))
    if _logo_ok:
        try:
            run = bp.add_run()
            run.add_picture(str(effective_logo), width=Inches(2.2))
        except Exception:
            _logo_ok = False

    if not _logo_ok:
        styled_run(bp, "wipro",                 bold=True, color_hex=WIPRO_RED,  size_pt=22)
        styled_run(bp, ":",                      bold=True, color_hex=WIPRO_RED,  size_pt=22)
        styled_run(bp, "  healthplan services",            color_hex=WIPRO_NAVY, size_pt=20)

    # — Metadata row helper ────────────────────────────────────────────────────
    def _meta(row_idx, ll, lv, rl=None, rv=None, merge_right=False):
        row = top.rows[row_idx]
        c0, c1, c2, c3 = row.cells
        _style_cell(c0, GRAY_LABEL); _col_width(c0, CL);  _label_para(c0, ll)
        if merge_right:
            c1 = c1.merge(c3)
            _style_cell(c1, WHITE); _col_width(c1, CM + CRL + CRV); _value_para(c1, lv)
        else:
            _style_cell(c1, WHITE);     _col_width(c1, CM);  _value_para(c1, lv)
            _style_cell(c2, GRAY_LABEL); _col_width(c2, CRL); _label_para(c2, rl or "")
            _style_cell(c3, WHITE);     _col_width(c3, CRV); _value_para(c3, rv or "")
        _no_row_break(row)

    _meta(1,  "Policy Name",          data.get("policy_name",  ""),   merge_right=True)
    _meta(2,  "Policy Number",        data.get("policy_number",""),
              "Version Number",        data.get("version",      ""))

    # Row 3 — GRC ID (blank left label/value)
    r3 = top.rows[3]
    for i, (shade, w) in enumerate([(GRAY_LABEL, CL), (WHITE, CM), (GRAY_LABEL, CRL), (WHITE, CRV)]):
        _style_cell(r3.cells[i], shade); _col_width(r3.cells[i], w)
        if i == 0: r3.cells[0].text = ""
        if i == 1: r3.cells[1].text = ""
        if i == 2: _label_para(r3.cells[2], "GRC ID Number")
        if i == 3: _value_para(r3.cells[3], data.get("grc_id", ""))
    _no_row_break(r3)

    _meta(4,  "Supersedes Policy",    data.get("supersedes",    ""),
              "Effective Date",        data.get("effective_date",""))
    _meta(5,  "Last Reviewed Date",   data.get("last_reviewed", ""),
              "Last Revised Date",     data.get("last_revised",  ""))
    _meta(6,  "Policy Custodian\nName(s)", data.get("custodians", ""), merge_right=True)

    # Row 7 — Owner / Approver subheaders
    r7 = top.rows[7]
    r7.cells[0].merge(r7.cells[1]); r7.cells[2].merge(r7.cells[3])
    _style_cell(r7.cells[0], GRAY_SUBHDR); _center_bold_para(r7.cells[0], "Policy Owner")
    _style_cell(r7.cells[2], GRAY_SUBHDR); _center_bold_para(r7.cells[2], "Policy Approver")
    _no_row_break(r7)

    _meta(8,  "Name",  data.get("owner_name",    ""), "Name",  data.get("approver_name",  ""))
    _meta(9,  "Title", data.get("owner_title",   ""), "Title", data.get("approver_title", ""))

    # Row 10 — Signature row
    r10 = top.rows[10]
    for i, (shade, w, lbl) in enumerate(
        [(GRAY_LABEL, CL, "Signature"), (WHITE, CM, None),
         (GRAY_LABEL, CRL, "Signature"), (WHITE, CRV, None)]
    ):
        c = r10.cells[i]; _style_cell(c, shade); _col_width(c, w)
        if lbl: _label_para(c, lbl)
        else:   c.text = ""
    _row_height(r10, 480)   # atLeast 0.33"
    _no_row_break(r10)

    _meta(11, "Date Signed",  data.get("date_signed",  ""),
              "Date Approved", data.get("date_approved",""))

    _gap()

    # ──────────────────────────────────────────────────────────────────────────
    # APPLICABLE TO TABLE
    # ──────────────────────────────────────────────────────────────────────────
    app_to  = data.get("applicable_to",    {})
    pol_typ = data.get("policy_types",     {})
    lob     = data.get("line_of_business", {})

    rows_def = [
        ("HealthPlan Services, Inc.",                      app_to.get("hps_inc",     False)),
        ("HealthPlan Services Insurance Agency, LLC",      app_to.get("agency",      True)),
        ("Policy Types",                                   None),   # sub-header
        ("Corporate",                                      app_to.get("corporate",   True)),
        ("Government Affairs Review Required",             app_to.get("govt_affairs",False)),
        ("Legal Review Required",                          app_to.get("legal_review",False)),
        ("Line of Business (LOB)",                         None),   # sub-header
        ("All LOBs",                                       lob.get("all_lobs",           True)),
        (f"Specific LOB: {lob.get('specific_lob','') or '[INSERT HERE]'}",
                                                           lob.get("specific_lob_checked",False)),
    ]

    AL = int(W * 0.28)    # left label column
    AR = W - AL

    t2 = _new_table(doc, len(rows_def), 2, [AL, AR], W)

    # Merge entire left column
    left_anchor = t2.rows[0].cells[0]
    for i in range(1, len(rows_def)):
        left_anchor = left_anchor.merge(t2.rows[i].cells[0])

    _style_cell(left_anchor, GRAY_LABEL)
    _cell_margins(left_anchor, top=160, bottom=160, left=120, right=120)
    _cell_valign(left_anchor, WD_ALIGN_VERTICAL.CENTER)
    left_anchor.text = ""
    lp = left_anchor.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(lp, 0, 0)
    styled_run(lp, "Applicable To:\n(select all that apply)", bold=True, size_pt=10.0)

    for i, (label, checked) in enumerate(rows_def):
        rc = t2.rows[i].cells[1]
        is_hdr = checked is None
        _style_cell(rc, GRAY_SUBHDR if is_hdr else WHITE)
        _cell_margins(rc, top=70, bottom=70, left=90, right=90)
        _col_width(rc, AR)
        _cell_valign(rc, WD_ALIGN_VERTICAL.CENTER)
        rc.text = ""
        p = rc.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _para_spacing(p, 20, 20)
        if is_hdr:
            styled_run(p, label, bold=True, size_pt=9.5)
        else:
            mark = "\u2611" if checked else "\u2610"
            styled_run(p, f"{label}  {mark}", size_pt=9.25)
        _no_row_break(t2.rows[i])

    _gap()

    # ──────────────────────────────────────────────────────────────────────────
    # SECTION TABLE HELPER
    # Key fix: content row does NOT use _no_row_break — it can flow across pages.
    # Header row keeps _keep_with_next so it travels with its first content line.
    # ──────────────────────────────────────────────────────────────────────────
    def _section(heading, fill_fn):
        tbl = _new_table(doc, 2, 1, [W], W)

        hdr = tbl.rows[0].cells[0]
        _style_cell(hdr, GRAY_SECTION)
        _cell_margins(hdr, top=45, bottom=45, left=80, right=80)
        _section_hdr_para(hdr, heading)
        _no_row_break(tbl.rows[0])     # header row: never orphaned

        cnt = tbl.rows[1].cells[0]
        _style_cell(cnt, WHITE)
        _cell_margins(cnt, top=60, bottom=80, left=120, right=120)
        cnt.text = ""
        fill_fn(cnt)
        # NOTE: no _no_row_break on content row — long content flows freely

        return tbl

    # ── Purpose ───────────────────────────────────────────────────────────────
    def _purpose(cell):
        for line in data.get("purpose", "").strip().split("\n"):
            line = line.strip()
            if line:
                _content_para(cell, line)
            else:
                _empty_para(cell)

    _section("Purpose and Scope", _purpose)
    _gap()

    # ── Definitions ───────────────────────────────────────────────────────────
    def _definitions(cell):
        defs = data.get("definitions") or {}
        if not defs:
            _content_para(cell, "")
            return
        for term, defn in defs.items():
            para = cell.add_paragraph()
            _para_spacing(para, 30, 30)
            styled_run(para, "\u2013  ",            size_pt=10)
            styled_run(para, f"{term}:  ",  bold=True, size_pt=10)
            styled_run(para, str(defn),     size_pt=10)

    _section("Definitions", _definitions)
    _gap()

    # ── Policy Statement ──────────────────────────────────────────────────────
    def _policy_statement(cell):
        stmt = data.get("policy_statement", "") or ""
        para = cell.add_paragraph()
        _para_spacing(para, 40, 40)
        # Style the opening clause ("It is the policy of ... that") in bold-italic
        lower = stmt.lower()
        idx = lower.find(" that ") + len(" that ") if " that " in lower else 0
        if idx:
            styled_run(para, stmt[:idx], bold=True, italic=True, size_pt=10)
            styled_run(para, stmt[idx:], size_pt=10)
        else:
            styled_run(para, stmt, size_pt=10)

    _section("Policy Statement", _policy_statement)
    _gap()

    # ── Procedures ────────────────────────────────────────────────────────────
    def _procedures(cell):
        for item in data.get("procedures", []):
            kind = item.get("type", "para")
            text = item.get("text", "")

            if kind == "empty":
                _empty_para(cell)
            elif kind == "heading":
                _heading_para(cell, text)
            elif kind == "bullet":
                _bullet_para(cell, text, is_sub=False)
            elif kind == "sub-bullet":
                _bullet_para(cell, text, is_sub=True)
            elif kind == "bold_intro":
                _content_para(cell, "", bold_prefix=(item.get("bold",""), item.get("rest","")))
            elif kind == "bold_intro_semi":
                para = cell.add_paragraph()
                _para_spacing(para, 40, 30)
                styled_run(para, item.get("bold", ""), bold=True, size_pt=10)
                segs = [s.strip() for s in item.get("rest","").split(";") if s.strip()]
                for i, seg in enumerate(segs):
                    suffix = ";" if i < len(segs) - 1 else ""
                    if i == 0:
                        styled_run(para, seg + suffix, size_pt=10)
                    else:
                        p2 = cell.add_paragraph()
                        _para_spacing(p2, 0, 30)
                        styled_run(p2, seg + suffix, size_pt=10)
            else:  # "para"
                if ";" in text:
                    _semi_breaks(cell, text)
                else:
                    _content_para(cell, text)

    _section("Procedures", _procedures)
    _gap()

    # ── Related Policies ──────────────────────────────────────────────────────
    def _related(cell):
        items = data.get("related_policies") or []
        if not items:
            _content_para(cell, "")
            return
        for pol in items:
            _content_para(cell, str(pol))

    _section("Related Policies or Standard Operating Procedures", _related)
    _gap()

    # ── Citations ─────────────────────────────────────────────────────────────
    def _citations(cell):
        items = data.get("citations") or []
        if not items:
            _content_para(cell, "")
            return
        for cit in items:
            if ";" in str(cit):
                _semi_breaks(cell, str(cit))
            else:
                _content_para(cell, str(cit))

    _section("Citations/References", _citations)
    _gap()

    # ──────────────────────────────────────────────────────────────────────────
    # REVISION HISTORY TABLE
    # ──────────────────────────────────────────────────────────────────────────
    rev_raw     = data.get("revision_history") or []
    rev_entries = []
    for entry in rev_raw:
        try:
            rev_entries.append(_norm_revision(entry))
        except Exception:
            pass   # skip malformed entries silently

    RC1 = int(W * 0.12)
    RC2 = int(W * 0.13)
    RC3 = int(W * 0.20)
    RC4 = W - RC1 - RC2 - RC3

    total_rev_rows = 2 + len(rev_entries)
    rt = _new_table(doc, total_rev_rows, 4, [RC1, RC2, RC3, RC4], W)

    # Heading span
    rh = rt.rows[0].cells[0]
    rh = rh.merge(rt.rows[0].cells[3])
    _style_cell(rh, GRAY_SECTION)
    _cell_margins(rh, top=45, bottom=45, left=80, right=80)
    _section_hdr_para(rh, "Revision History")
    _no_row_break(rt.rows[0])

    # Column headers
    for ci, (hdr_text, w) in enumerate([
        ("Date", RC1), ("Version Number", RC2),
        ("Updated By", RC3), ("Description of Update", RC4)
    ]):
        hc = rt.rows[1].cells[ci]
        _style_cell(hc, GRAY_LABEL); _col_width(hc, w)
        _value_para(hc, hdr_text, size_pt=9.0, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    _no_row_break(rt.rows[1])

    # Data rows
    for ri, entry in enumerate(rev_entries, start=2):
        row = rt.rows[ri]
        for ci, (txt, w) in enumerate(zip(entry, [RC1, RC2, RC3, RC4])):
            rc = row.cells[ci]
            _style_cell(rc, WHITE); _col_width(rc, w)
            rc.text = ""
            lines = str(txt).split("\n")
            for li, line in enumerate(lines):
                para = rc.paragraphs[0] if li == 0 else rc.add_paragraph()
                _para_spacing(para, 20, 20)
                styled_run(para, line.strip(), size_pt=9.0)
        # Content rows flow freely — no _no_row_break

    # ──────────────────────────────────────────────────────────────────────────
    # FOOTER
    # ──────────────────────────────────────────────────────────────────────────
    footer     = sec.footer
    footer_tbl = footer.add_table(rows=2, cols=1, width=Inches(7.0))
    footer_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r in footer_tbl.rows:
        _no_row_break(r)
        c = r.cells[0]
        c.text = ""
        _cell_borders(c, color=WHITE, size=0)
        _cell_shade(c, WHITE)
        _cell_margins(c, top=0, bottom=0, left=0, right=0)
        _cell_valign(c, WD_ALIGN_VERTICAL.CENTER)

    fp1 = footer_tbl.rows[0].cells[0].paragraphs[0]
    fp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(fp1, 0, 0)
    styled_run(
        fp1,
        "Confidential & Proprietary \u00A9 HealthPlan Services Inc.,"
        " including its subsidiaries and affiliates",
        size_pt=7.5, color_hex=FOOTER_GRAY,
    )

    fp2 = footer_tbl.rows[1].cells[0].paragraphs[0]
    fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(fp2, 0, 0)
    styled_run(
        fp2,
        f"{data.get('policy_number','')}  |  "
        f"{data.get('policy_name','')}  |  "
        f"{data.get('version','')}",
        size_pt=7.5, color_hex=FOOTER_GRAY,
    )

    doc.save(output_path)
    return output_path
