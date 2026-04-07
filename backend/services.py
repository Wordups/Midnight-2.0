"""
services.py — Midnight backend processing layer

Pipeline:
    get_uploaded_text()
        ↓
    run_llm_transform()       → POLICY_DATA
        ↓
    run_framework_mapping()   → framework_map{}
        ↓
    build_output_doc()        → policy.docx
    build_grc_summary_doc()   → grc_summary.pdf
"""

from __future__ import annotations

import io
import json
import os
import re
import tempfile
import uuid
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from groq import Groq

from hps_policy_migration_builder import build_policy_document


# ── Template registry ─────────────────────────────────────────────────────────

try:
    from templates.template_generic import build_document as _build_generic
    _TEMPLATES_AVAILABLE = True
except ImportError:
    _TEMPLATES_AVAILABLE = False
    _build_generic = None

TEMPLATE_REGISTRY = {
    "Generic Policy Template": "_generic",
    "Generic": "_generic",
    "Enterprise Policy Template": "_generic",
    "Enterprise Template A": "_hps",
}


# ── Runtime config ────────────────────────────────────────────────────────────

GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
GROQ_MODEL = os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile")
MIN_TEXT_LEN = 50
MAX_SOURCE_LEN = 24000
MAX_POLICY_ANALYSIS_LEN = 16000


# ── Prompts ───────────────────────────────────────────────────────────────────

EXTRACTION_PROMPT = """
You are a policy migration specialist.

Read the attached legacy policy document and extract ALL content into the exact
JSON structure below.

STRICT RULES:
- Do NOT summarize, rewrite, or remove content
- Preserve source wording as closely as possible
- Fix only minor spacing / punctuation / obvious grammar issues
- Map all content into the correct field
- If a field is missing in the source, return an empty string, empty object, or empty list
- Preserve revision history if present
- Preserve related policies and citations if present
- Preserve definitions if present
- Preserve procedures in the correct order
- For procedure items use exactly one type:
  "para"            = standalone paragraph
  "heading"         = bold underlined subsection title
  "bullet"          = first-level bullet
  "sub-bullet"      = second-level bullet
  "bold_intro"      = paragraph starting with bold label; keys: "bold" and "rest"
  "bold_intro_semi" = same as bold_intro but "rest" contains semicolons
  "empty"           = blank spacer line

OUTPUT RULES:
- Return ONLY valid JSON
- No explanation
- No markdown fences
- No comments
- No preamble
- No trailing commas
- Use double quotes for all keys and string values
- Escape quotes inside strings with backslash
- Start with {
- End with }

JSON SCHEMA:
{
  "policy_name": "",
  "policy_number": "",
  "version": "",
  "grc_id": "",
  "supersedes": "",
  "effective_date": "",
  "last_reviewed": "",
  "last_revised": "",
  "custodians": "",
  "owner_name": "",
  "owner_title": "",
  "approver_name": "",
  "approver_title": "",
  "date_signed": "",
  "date_approved": "",
  "applicable_to": {
    "hps_inc": false,
    "agency": true,
    "corporate": true,
    "govt_affairs": false,
    "legal_review": false
  },
  "policy_types": {
    "carrier_specific": false,
    "cross_carrier": false,
    "global": false,
    "on_off_hix": false
  },
  "line_of_business": {
    "all_lobs": true,
    "specific_lob": "",
    "specific_lob_checked": false
  },
  "purpose": "",
  "definitions": {},
  "policy_statement": "",
  "procedures": [],
  "related_policies": [],
  "citations": [],
  "revision_history": []
}

REVISION HISTORY FORMAT:
[
  {
    "date": "",
    "version": "",
    "updated_by": "",
    "description": ""
  }
]

PROCEDURES FORMAT:
[
  {"type": "para",            "text": ""},
  {"type": "heading",         "text": ""},
  {"type": "bullet",          "text": ""},
  {"type": "sub-bullet",      "text": ""},
  {"type": "bold_intro",      "bold": "", "rest": ""},
  {"type": "bold_intro_semi", "bold": "", "rest": ""},
  {"type": "empty"}
]

HERE IS THE LEGACY POLICY DOCUMENT:
"""

JSON_REPAIR_PROMPT = """
You are repairing malformed JSON from a policy extraction pipeline.

Your job:
- Convert the input into valid JSON only
- Preserve ALL content exactly — do NOT summarize or remove fields
- Escape inner quotes correctly using backslash
- Fix commas, brackets, and broken strings only
- Return only one valid JSON object
- No explanation, no markdown, no preamble
- Start with {
- End with }

BROKEN JSON:
"""

FRAMEWORK_MAPPING_PROMPT = """
You are a compliance framework specialist with deep expertise in:
- HIPAA (45 CFR 164)
- HiTrust CSF v9.3 and v11
- PCI DSS v3.2 and v4.0
- ISO/IEC 27001:2022
- CoBIT 5.0 and 2019
- NIST Cybersecurity Framework v1.1 and v2.0
- NIST SP 800-53
- SOC 2 Type II

Analyze this policy and produce a complete framework mapping.

Do FOUR things:

1. MAPPED CITATIONS — for each policy section that satisfies a control,
   cite the exact control number. Be specific.

2. GAP ANALYSIS — identify controls this policy SHOULD address based on
   its type and subject matter but does not. Only flag genuine relevant gaps.

3. SUGGESTIONS — for each gap, write 1-3 sentences of specific policy language
   the organization could add to close it. Write in formal policy language.

4. AUDIT SUMMARY — 2-3 plain-language sentences a non-technical person can read
   to understand: what frameworks are covered, how many gaps exist, overall posture.

Return ONLY valid JSON.
- No explanation, no markdown, no preamble, no comments, no trailing commas
- Start with { and end with }

{
  "policy_name": "",
  "policy_type": "",
  "overall_coverage": "strong|moderate|weak|unknown",
  "mapped_citations": [
    {
      "framework": "",
      "control_id": "",
      "control_name": "",
      "policy_section": "",
      "coverage_note": ""
    }
  ],
  "gaps": [
    {
      "framework": "",
      "control_id": "",
      "control_name": "",
      "gap_description": "",
      "risk_level": "high|medium|low",
      "suggestion": ""
    }
  ],
  "audit_summary": "",
  "frameworks_covered": [],
  "total_controls_mapped": 0,
  "total_gaps": 0
}

HERE IS THE POLICY TO ANALYZE:
"""


# ── Generic text sanitation ───────────────────────────────────────────────────

def _remove_illegal_control_chars(text: str, preserve_newlines: bool = True) -> str:
    if not text:
        return ""

    allowed = "\n\t" if preserve_newlines else ""
    bad_controls = "".join(chr(i) for i in range(32) if chr(i) not in allowed)
    return text.translate(str.maketrans("", "", bad_controls))


def _normalize_unicode_punctuation(text: str) -> str:
    if not text:
        return ""

    replacements = {
        "\u201c": '"',
        "\u201d": '"',
        "\u2018": "'",
        "\u2019": "'",
        "\u2013": "-",
        "\u2014": "-",
        "\u2022": "•",
        "\u25e6": "◦",
        "\u00a0": " ",
        "\u200b": "",
        "\ufeff": "",
        "\u2028": " ",
        "\u2029": " ",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    return text


def _clean_text(text: str, preserve_newlines: bool = True) -> str:
    if text is None:
        return ""

    text = str(text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _remove_illegal_control_chars(text, preserve_newlines=preserve_newlines)
    text = _normalize_unicode_punctuation(text)

    if preserve_newlines:
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = "\n".join(line.strip() for line in text.splitlines())
    else:
        text = re.sub(r"\s+", " ", text)

    return text.strip()


def _clean_scalar(value: Any) -> str:
    return _clean_text(str(value or ""), preserve_newlines=False)


def _safe_int(value: Any, default: int = 0) -> int:
    try:
        if value in (None, "", False):
            return default
        return int(value)
    except (TypeError, ValueError):
        return default


def _dedupe_preserve_order(items: list[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for item in items:
        cleaned = _clean_scalar(item)
        if cleaned and cleaned not in seen:
            seen.add(cleaned)
            result.append(cleaned)
    return result


# ── Extraction helpers ────────────────────────────────────────────────────────

def _extract_docx_bytes(file_bytes: bytes) -> str:
    lines: list[str] = []

    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(f"Could not open .docx file: {e}")

    for para in doc.paragraphs:
        text = _clean_text(para.text, preserve_newlines=False)
        if text:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_parts: list[str] = []
            for cell in row.cells:
                cell_text = "\n".join(
                    _clean_text(p.text, preserve_newlines=False)
                    for p in cell.paragraphs
                    if _clean_text(p.text, preserve_newlines=False)
                ).strip()
                if cell_text:
                    row_parts.append(cell_text)
            if row_parts:
                lines.append(" | ".join(_dedupe_preserve_order(row_parts)))

    for section in doc.sections:
        try:
            for hdr_para in section.header.paragraphs:
                text = _clean_text(hdr_para.text, preserve_newlines=False)
                if text:
                    lines.append(f"[HEADER] {text}")
        except Exception:
            pass

        try:
            for ftr_para in section.footer.paragraphs:
                text = _clean_text(ftr_para.text, preserve_newlines=False)
                if text:
                    lines.append(f"[FOOTER] {text}")
        except Exception:
            pass

    return _clean_text("\n".join(lines), preserve_newlines=True)


def _extract_txt_bytes(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            return _clean_text(file_bytes.decode(encoding), preserve_newlines=True)
        except UnicodeDecodeError:
            continue
    return _clean_text(file_bytes.decode("utf-8", errors="replace"), preserve_newlines=True)


def get_uploaded_text(filename: str, file_bytes: bytes) -> str:
    name = filename.lower().strip()

    if name.endswith(".docx"):
        text = _extract_docx_bytes(file_bytes)
    elif name.endswith((".txt", ".md")):
        text = _extract_txt_bytes(file_bytes)
    else:
        try:
            text = _extract_docx_bytes(file_bytes)
        except Exception:
            text = _extract_txt_bytes(file_bytes)

    text = _clean_text(text, preserve_newlines=True)

    print(f"\n{'=' * 60}\nEXTRACTION  |  {filename}  |  {len(text)} chars\n{'=' * 60}\n")
    return text


# ── LLM helpers ───────────────────────────────────────────────────────────────

def _groq_call(prompt: str, content: str, max_tokens: int = 8000) -> str:
    if not GROQ_API_KEY:
        raise ValueError("GROQ_API_KEY not set.")

    client = Groq(api_key=GROQ_API_KEY)
    response = client.chat.completions.create(
        model=GROQ_MODEL,
        messages=[{"role": "user", "content": prompt + "\n\n" + content}],
        temperature=0.1,
        max_tokens=max_tokens,
    )

    return _clean_text(response.choices[0].message.content or "", preserve_newlines=True)


def _sanitize_llm_output(raw: str) -> str:
    """
    Sanitize LLM output before JSON parsing.
    This is the critical hardening point for malformed control characters,
    smart punctuation, zero-width chars, and broken line endings.
    """
    return _clean_text(raw, preserve_newlines=True)


def _strip_code_fences(raw: str) -> str:
    raw = raw.strip()

    if raw.startswith("```"):
        raw = re.sub(r"^```(?:json|python)?\s*", "", raw, flags=re.IGNORECASE)
        raw = re.sub(r"\s*```$", "", raw)

    return raw.strip()


def _extract_json_object(raw: str) -> str:
    raw = _sanitize_llm_output(raw)
    raw = _strip_code_fences(raw)

    if raw.startswith("POLICY_DATA ="):
        raw = raw.split("=", 1)[1].strip()
    elif raw.startswith("POLICY_DATA="):
        raw = raw.split("=", 1)[1].strip()

    first = raw.find("{")
    last = raw.rfind("}")

    if first == -1 or last == -1 or last < first:
        raise ValueError("Model response did not contain a JSON object.")

    raw_json = raw[first:last + 1]
    return _sanitize_llm_output(raw_json)


def _json_error_context(raw: str, lineno: int | None, window: int = 6) -> str:
    lines = raw.splitlines()
    if not lines:
        return raw[:600]

    line_no = lineno or 1
    start = max(0, line_no - window - 1)
    end = min(len(lines), line_no + window)

    return "\n".join(f"{i + 1}: {lines[i]}" for i in range(start, end))


def _repair_json_via_llm(raw_json: str) -> str:
    repaired = _groq_call(JSON_REPAIR_PROMPT, raw_json, max_tokens=8000)
    return _extract_json_object(repaired)


def _attempt_json_parse(raw_json: str) -> dict[str, Any]:
    return json.loads(raw_json)


# ── Normalizers ───────────────────────────────────────────────────────────────

def _normalize_revision_history(value: Any) -> list[dict[str, str]]:
    if not value:
        return []

    out: list[dict[str, str]] = []

    if isinstance(value, list):
        for entry in value:
            if isinstance(entry, dict):
                out.append({
                    "date": _clean_scalar(entry.get("date", "")),
                    "version": _clean_scalar(entry.get("version", "")),
                    "updated_by": _clean_scalar(entry.get("updated_by", "")),
                    "description": _clean_scalar(entry.get("description", "")),
                })
            elif isinstance(entry, (list, tuple)):
                padded = list(entry) + ["", "", "", ""]
                out.append({
                    "date": _clean_scalar(padded[0]),
                    "version": _clean_scalar(padded[1]),
                    "updated_by": _clean_scalar(padded[2]),
                    "description": _clean_scalar(padded[3]),
                })
            else:
                out.append({
                    "date": "",
                    "version": "",
                    "updated_by": "",
                    "description": _clean_scalar(entry),
                })

    return out


def _normalize_procedures(value: Any) -> list[dict[str, Any]]:
    allowed = {
        "para",
        "heading",
        "bullet",
        "sub-bullet",
        "bold_intro",
        "bold_intro_semi",
        "empty",
    }

    if not isinstance(value, list):
        return []

    normalized: list[dict[str, Any]] = []

    for item in value:
        if isinstance(item, str):
            text = _clean_scalar(item)
            if text:
                normalized.append({"type": "para", "text": text})
            continue

        if not isinstance(item, dict):
            continue

        kind = _clean_scalar(item.get("type", "para")).lower()
        if kind not in allowed:
            kind = "para"

        if kind in {"bold_intro", "bold_intro_semi"}:
            normalized.append({
                "type": kind,
                "bold": _clean_scalar(item.get("bold", "")),
                "rest": _clean_scalar(item.get("rest", "")),
            })
        elif kind == "empty":
            normalized.append({"type": "empty"})
        else:
            text = _clean_scalar(item.get("text", ""))
            if text:
                normalized.append({
                    "type": kind,
                    "text": text,
                })

    return normalized


def _coerce_bool(value: Any, default: bool = False) -> bool:
    if isinstance(value, bool):
        return value

    if isinstance(value, str):
        v = value.strip().lower()
        if v in {"true", "yes", "y", "1", "checked", "x"}:
            return True
        if v in {"false", "no", "n", "0", "unchecked"}:
            return False

    if isinstance(value, int):
        return bool(value)

    return default


def _normalize_string_list(value: Any) -> list[str]:
    if not isinstance(value, list):
        return []
    return _dedupe_preserve_order([str(x) for x in value if _clean_scalar(x)])


def _normalize_definitions(value: Any) -> dict[str, str]:
    if not isinstance(value, dict):
        return {}

    out: dict[str, str] = {}
    for k, v in value.items():
        key = _clean_scalar(k)
        val = _clean_scalar(v)
        if key:
            out[key] = val
    return out


def _validate_policy_data(data: dict[str, Any]) -> dict[str, Any]:
    if not isinstance(data, dict):
        raise ValueError("Parsed POLICY_DATA is not a JSON object.")

    scalar_fields = [
        "policy_name",
        "policy_number",
        "version",
        "grc_id",
        "supersedes",
        "effective_date",
        "last_reviewed",
        "last_revised",
        "custodians",
        "owner_name",
        "owner_title",
        "approver_name",
        "approver_title",
        "date_signed",
        "date_approved",
        "purpose",
        "policy_statement",
    ]

    for key in scalar_fields:
        data[key] = _clean_scalar(data.get(key, ""))

    app = data.get("applicable_to", {}) if isinstance(data.get("applicable_to"), dict) else {}
    data["applicable_to"] = {
        "hps_inc": _coerce_bool(app.get("hps_inc"), False),
        "agency": _coerce_bool(app.get("agency"), True),
        "corporate": _coerce_bool(app.get("corporate"), True),
        "govt_affairs": _coerce_bool(app.get("govt_affairs"), False),
        "legal_review": _coerce_bool(app.get("legal_review"), False),
    }

    pol = data.get("policy_types", {}) if isinstance(data.get("policy_types"), dict) else {}
    data["policy_types"] = {
        "carrier_specific": _coerce_bool(pol.get("carrier_specific"), False),
        "cross_carrier": _coerce_bool(pol.get("cross_carrier"), False),
        "global": _coerce_bool(pol.get("global"), False),
        "on_off_hix": _coerce_bool(pol.get("on_off_hix"), False),
    }

    lob = data.get("line_of_business", {}) if isinstance(data.get("line_of_business"), dict) else {}
    data["line_of_business"] = {
        "all_lobs": _coerce_bool(lob.get("all_lobs"), True),
        "specific_lob": _clean_scalar(lob.get("specific_lob", "")),
        "specific_lob_checked": _coerce_bool(lob.get("specific_lob_checked"), False),
    }

    data["definitions"] = _normalize_definitions(data.get("definitions", {}))
    data["related_policies"] = _normalize_string_list(data.get("related_policies", []))
    data["citations"] = _normalize_string_list(data.get("citations", []))
    data["procedures"] = _normalize_procedures(data.get("procedures", []))
    data["revision_history"] = _normalize_revision_history(data.get("revision_history", []))

    required = ["policy_name", "policy_number", "version"]
    missing = [k for k in required if not data.get(k)]
    if missing:
        raise ValueError(f"Missing required fields after extraction: {', '.join(missing)}")

    return data


def _parse_policy_data(raw: str) -> dict[str, Any]:
    raw_json = _extract_json_object(raw)

    try:
        parsed = _attempt_json_parse(raw_json)
    except json.JSONDecodeError as e:
        print(f"[WARN] Initial POLICY_DATA JSON parse failed at line {e.lineno}, col {e.colno}")
        print(_json_error_context(raw_json, e.lineno))

        # First retry after an additional sanitize pass
        try:
            raw_json_retry = _sanitize_llm_output(raw_json)
            parsed = _attempt_json_parse(raw_json_retry)
            print("[INFO] POLICY_DATA parsed successfully after secondary sanitization pass.")
        except Exception:
            # Final retry via LLM JSON repair
            try:
                repaired_json = _repair_json_via_llm(raw_json)
                parsed = _attempt_json_parse(repaired_json)
                print("[INFO] POLICY_DATA repaired successfully by secondary JSON repair pass.")
            except Exception as repair_error:
                context = _json_error_context(raw_json, e.lineno)
                raise ValueError(
                    f"JSON parsing failed at line {e.lineno}, column {e.colno}.\n\n"
                    f"Context:\n{context}\n\n"
                    f"Original error: {e}\n\n"
                    f"Repair attempt failed: {repair_error}"
                ) from repair_error

    return _validate_policy_data(parsed)


def run_llm_transform(source_text: str, template_name: str) -> dict[str, Any]:
    source_text = _clean_text(source_text, preserve_newlines=True)

    if len(source_text) < MIN_TEXT_LEN:
        raise ValueError(f"Extracted text too short ({len(source_text)} chars).")

    if len(source_text) > MAX_SOURCE_LEN:
        source_text = source_text[:MAX_SOURCE_LEN]

    try:
        raw = _groq_call(EXTRACTION_PROMPT, source_text, max_tokens=8000)
    except Exception as e:
        raise ValueError(f"Groq extraction failed: {e}") from e

    policy_data = _parse_policy_data(raw)
    policy_data["template_name"] = _clean_scalar(template_name) or "Generic Policy Template"
    return policy_data


# ── Framework Mapping ─────────────────────────────────────────────────────────

def _policy_to_text(policy_data: dict[str, Any]) -> str:
    lines: list[str] = []

    lines.append(
        f"POLICY: {policy_data.get('policy_name', '')} "
        f"({policy_data.get('policy_number', '')})"
    )
    lines.append(f"VERSION: {policy_data.get('version', '')}")
    lines.append("")

    if policy_data.get("purpose"):
        lines.append("PURPOSE:")
        lines.append(policy_data["purpose"])
        lines.append("")

    if policy_data.get("policy_statement"):
        lines.append("POLICY STATEMENT:")
        lines.append(policy_data["policy_statement"])
        lines.append("")

    defs = policy_data.get("definitions", {})
    if defs:
        lines.append("DEFINITIONS:")
        for k, v in defs.items():
            lines.append(f"  {k}: {v}")
        lines.append("")

    procs = policy_data.get("procedures", [])
    if procs:
        lines.append("PROCEDURES:")
        for item in procs:
            kind = item.get("type", "")
            if kind == "heading":
                lines.append(f"  [{item.get('text', '')}]")
            elif kind == "bullet":
                lines.append(f"  • {item.get('text', '')}")
            elif kind == "sub-bullet":
                lines.append(f"    ◦ {item.get('text', '')}")
            elif kind in ("bold_intro", "bold_intro_semi"):
                lines.append(f"  {item.get('bold', '')} {item.get('rest', '')}".strip())
            elif kind == "para":
                lines.append(f"  {item.get('text', '')}")
            elif kind == "empty":
                lines.append("")

        lines.append("")

    existing = policy_data.get("citations", [])
    if existing:
        lines.append("EXISTING CITATIONS:")
        for c in existing:
            lines.append(f"  {c}")

    return _clean_text("\n".join(lines), preserve_newlines=True)


def _empty_map(
    policy_data: dict[str, Any],
    reason: str = "Framework mapping unavailable.",
) -> dict[str, Any]:
    return {
        "policy_name": policy_data.get("policy_name", ""),
        "policy_type": "",
        "overall_coverage": "unknown",
        "mapped_citations": [],
        "gaps": [],
        "audit_summary": reason,
        "frameworks_covered": [],
        "total_controls_mapped": 0,
        "total_gaps": 0,
        "_mapping_failed": True,
        "_mapping_failure_reason": reason,
    }


def _normalize_framework_entries(value: Any) -> list[dict[str, str]]:
    if not isinstance(value, list):
        return []

    out: list[dict[str, str]] = []
    for entry in value:
        if not isinstance(entry, dict):
            continue

        cleaned = {
            str(k): _clean_scalar(v)
            for k, v in entry.items()
            if _clean_scalar(k)
        }
        if cleaned:
            out.append(cleaned)

    return out


def _parse_framework_map(raw: str, policy_data: dict[str, Any]) -> dict[str, Any]:
    raw = _sanitize_llm_output(raw)
    raw = _strip_code_fences(raw)

    first = raw.find("{")
    last = raw.rfind("}")

    if first == -1 or last == -1 or last < first:
        raise ValueError("Framework mapping response did not contain a valid JSON object.")

    raw_json = raw[first:last + 1]

    try:
        framework_map = json.loads(raw_json)
    except json.JSONDecodeError as e:
        context = _json_error_context(raw_json, e.lineno)
        raise ValueError(
            f"Framework JSON parse failed at line {e.lineno}, col {e.colno}.\n\nContext:\n{context}"
        ) from e

    if not isinstance(framework_map, dict):
        raise ValueError("Framework mapping response was not a JSON object.")

    framework_map["policy_name"] = _clean_scalar(
        framework_map.get("policy_name", policy_data.get("policy_name", ""))
    )
    framework_map["policy_type"] = _clean_scalar(framework_map.get("policy_type", ""))

    coverage = _clean_scalar(framework_map.get("overall_coverage", "unknown")).lower()
    framework_map["overall_coverage"] = (
        coverage if coverage in {"strong", "moderate", "weak", "unknown"} else "unknown"
    )

    framework_map["mapped_citations"] = _normalize_framework_entries(
        framework_map.get("mapped_citations", [])
    )
    framework_map["gaps"] = _normalize_framework_entries(framework_map.get("gaps", []))
    framework_map["frameworks_covered"] = _normalize_string_list(
        framework_map.get("frameworks_covered", [])
    )
    framework_map["audit_summary"] = _clean_scalar(framework_map.get("audit_summary", ""))

    framework_map["total_controls_mapped"] = _safe_int(
        framework_map.get("total_controls_mapped"),
        default=len(framework_map["mapped_citations"]),
    )
    framework_map["total_gaps"] = _safe_int(
        framework_map.get("total_gaps"),
        default=len(framework_map["gaps"]),
    )

    if not framework_map["audit_summary"]:
        raise ValueError("Framework mapping missing audit_summary.")

    framework_map["_mapping_failed"] = False
    framework_map["_mapping_failure_reason"] = ""

    return framework_map


def run_framework_mapping(policy_data: dict[str, Any]) -> dict[str, Any]:
    policy_text = _policy_to_text(policy_data)

    if len(policy_text) > MAX_POLICY_ANALYSIS_LEN:
        policy_text = policy_text[:MAX_POLICY_ANALYSIS_LEN]

    print(
        f"\n{'=' * 60}\nFRAMEWORK   |  "
        f"{policy_data.get('policy_name', '')}  |  {len(policy_text)} chars\n{'=' * 60}\n"
    )

    try:
        raw = _groq_call(FRAMEWORK_MAPPING_PROMPT, policy_text, max_tokens=6000)

        print("\n--- RAW FRAMEWORK OUTPUT (START) ---\n")
        print(raw[:4000])
        print("\n--- RAW FRAMEWORK OUTPUT (END) ---\n")

        framework_map = _parse_framework_map(raw, policy_data)

    except Exception as e:
        print(f"[WARN] Framework mapping failed: {e}")
        framework_map = _empty_map(
            policy_data,
            reason=f"Framework mapping failed: {str(e)}"
        )

    print(
        "FRAMEWORK   |  "
        f"mapped: {framework_map.get('total_controls_mapped', 0)}  "
        f"gaps: {framework_map.get('total_gaps', 0)}  "
        f"coverage: {framework_map.get('overall_coverage', '?')}  "
        f"failed: {framework_map.get('_mapping_failed', False)}\n"
    )
    return framework_map


# ── GRC Summary Doc — PDF ────────────────────────────────────────────────────

def build_grc_summary_doc(
    policy_data: dict[str, Any],
    framework_map: dict[str, Any],
) -> tuple[str, bytes]:
    name = _clean_scalar(policy_data.get("policy_name", "Policy")) or "Policy"
    number = _clean_scalar(policy_data.get("policy_number", "SEC-P")) or "SEC-P"
    ver = _clean_scalar(policy_data.get("version", "V1.0")) or "V1.0"
    fname = f"{number} {name} {ver}-GRC-Summary.pdf"

    print(f"GRC PDF  |  {name}  |  {number}  |  {ver}")

    try:
        from grc_summary_pdf import build_grc_pdf
    except Exception as e:
        raise RuntimeError(f"[IMPORT ERROR] grc_summary_pdf failed: {e}") from e

    try:
        pdf_bytes = build_grc_pdf(policy_data, framework_map)
    except Exception as e:
        raise RuntimeError(f"[PDF ERROR] GRC generation failed: {e}") from e

    if not isinstance(pdf_bytes, (bytes, bytearray)) or not pdf_bytes:
        raise RuntimeError("[PDF ERROR] GRC generation returned empty or invalid PDF bytes.")

    return fname, bytes(pdf_bytes)


# ── Logo ──────────────────────────────────────────────────────────────────────

def save_logo_bytes(filename: str, file_bytes: bytes) -> str:
    tmp_dir = os.path.join(tempfile.gettempdir(), "midnight_logos")
    os.makedirs(tmp_dir, exist_ok=True)

    ext = Path(filename).suffix.lower() or ".png"
    stem = "".join(c for c in Path(filename).stem if c.isalnum() or c in "-_") or "logo"
    path = os.path.join(tmp_dir, f"{stem}_{uuid.uuid4().hex[:8]}{ext}")

    with open(path, "wb") as f:
        f.write(file_bytes)

    return path


# ── Policy Doc Builder ────────────────────────────────────────────────────────

def build_output_doc(
    policy_data: dict[str, Any],
    logo_path: str | None = None,
) -> tuple[str, bytes]:
    name = _clean_scalar(policy_data.get("policy_name", "Policy")) or "Policy"
    number = _clean_scalar(policy_data.get("policy_number", "POL")) or "POL"
    ver = _clean_scalar(policy_data.get("version", "V1.0")) or "V1.0"
    template_name = _clean_scalar(policy_data.get("template_name", "Generic Policy Template"))
    logo_pos = _clean_scalar(policy_data.get("logo_position", "left")).lower() or "left"
    fname = f"{number} {name} {ver}-NEW.docx"

    if logo_pos not in {"left", "center", "right"}:
        logo_pos = "left"

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp_path = tmp.name

    try:
        registry_key = TEMPLATE_REGISTRY.get(template_name)

        if registry_key == "_hps":
            print(f"BUILD  |  template: {template_name}  →  HPS builder")
            build_policy_document(policy_data, tmp_path, logo_path=logo_path)

        elif registry_key == "_generic" and _TEMPLATES_AVAILABLE and _build_generic:
            print(f"BUILD  |  template: {template_name}  →  generic builder")
            _build_generic(
                policy_data,
                tmp_path,
                logo_path=logo_path,
                logo_position=logo_pos,
            )

        elif registry_key is None:
            print(f"[WARN] Template '{template_name}' not in registry — falling back to generic")
            if _TEMPLATES_AVAILABLE and _build_generic:
                _build_generic(
                    policy_data,
                    tmp_path,
                    logo_path=logo_path,
                    logo_position=logo_pos,
                )
            else:
                print("[WARN] Generic builder unavailable — falling back to HPS")
                build_policy_document(policy_data, tmp_path, logo_path=logo_path)

        else:
            print(f"[WARN] Builder for '{registry_key}' not yet implemented — using generic")
            if _TEMPLATES_AVAILABLE and _build_generic:
                _build_generic(
                    policy_data,
                    tmp_path,
                    logo_path=logo_path,
                    logo_position=logo_pos,
                )
            else:
                build_policy_document(policy_data, tmp_path, logo_path=logo_path)

        with open(tmp_path, "rb") as f:
            doc_bytes = f.read()

    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass

    if not doc_bytes:
        raise RuntimeError("Generated DOCX was empty.")

    print(f"BUILD  |  template: {template_name}  |  file: {fname}")
    return fname, doc_bytes
