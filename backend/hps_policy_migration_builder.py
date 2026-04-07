"""
hps_policy_migration_builder.py

HPS Policy Migration Builder — Strict Template Renderer

Design goals:
- Preserve HPS template fidelity
- Keep rigid enterprise matrix structures intact
- Do not flatten checkbox rails into text
- Use deterministic fixed-layout tables
- Keep signature fields blank for review
- Allow longer metadata values (like custodians) without clipping
"""

from __future__ import annotations

import os
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches


# ──────────────────────────────────────────────────────────────────────────────
# PAGE GEOMETRY
# ──────────────────────────────────────────────────────────────────────────────

_PAGE_W_IN = 8.5
_PAGE_H_IN = 11.0
_MARGIN_LR_IN = 0.75
_MARGIN_TOP_IN = 0.75
_MARGIN_BOTTOM_IN = 0.90
_CONTENT_IN = _PAGE_W_IN - (_MARGIN_LR_IN * 2)   # 7.0"
_CONTENT_W = int(_CONTENT_IN * 1440)             # 10080 twips


# ──────────────────────────────────────────────────────────────────────────────
# COLORS / BRAND
# ──────────────────────────────────────────────────────────────────────────────

GRAY_BANNER = "D9D9D9"
GRAY_LABEL = "D9D9D9"
GRAY_SUBHDR = "BFBFBF"
GRAY_SECTION = "D9D9D9"
WHITE = "FFFFFF"
BLACK = "000000"
WIPRO_RED = "C00000"
WIPRO_NAVY = "17375E"
FOOTER_GRAY = "595959"


# ──────────────────────────────────────────────────────────────────────────────
# TYPOGRAPHY
# ──────────────────────────────────────────────────────────────────────────────

FONT_FAMILY = "Arial"
BASE_PT = 9.5
BODY_PT = 10.0
SMALL_PT = 7.5
LABEL_PT = 9.0

# Important:
# This controls the LOGO size, not the banner cell size.
BANNER_LOGO_WIDTH_IN = 3.35


# ──────────────────────────────────────────────────────────────────────────────
# GENERAL HELPERS
# ──────────────────────────────────────────────────────────────────────────────

def _safe(value: Any) -> str:
    return str(value or "").strip()


def _bool(value: Any) -> bool:
    return bool(value is True)


def _dict(value: Any) -> dict[str, Any]:
    return value if isinstance(value, dict) else {}


def _list(value: Any) -> list[Any]:
    return value if isinstance(value, list) else []


def _rgb(hex_str: str):
    h = hex_str.lstrip("#")
    return tuple(int(h[i:i + 2], 16) for i in (0, 2, 4))


def _remove(parent, tag):
    for old in parent.findall(qn(tag)):
        parent.remove(old)


# ──────────────────────────────────────────────────────────────────────────────
# LOW-LEVEL XML HELPERS
# ──────────────────────────────────────────────────────────────────────────────

def styled_run(
    para,
    text,
    *,
    bold=False,
    italic=False,
    underline=False,
    color_hex=BLACK,
    size_pt=BASE_PT,
    font=FONT_FAMILY,
):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.name = font
    run.font.size = Pt(size_pt)
    r, g, b = _rgb(color_hex)
    run.font.color.rgb = RGBColor(r, g, b)
    return run


def _para_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    _remove(pPr, "w:spacing")
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:before"), str(before))
    spc.set(qn("w:after"), str(after))
    if line is not None:
        spc.set(qn("w:line"), str(line))
        spc.set(qn("w:lineRule"), "auto")
    pPr.append(spc)


def _keep_with_next(para, on=True):
    pPr = para._p.get_or_add_pPr()
    _remove(pPr, "w:keepNext")
    if on:
        pPr.append(OxmlElement("w:keepNext"))


def _keep_lines(para, on=True):
    pPr = para._p.get_or_add_pPr()
    _remove(pPr, "w:keepLines")
    if on:
        pPr.append(OxmlElement("w:keepLines"))


def _cell_shade(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    _remove(tcPr, "w:shd")
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex.upper())
    tcPr.append(shd)


def _cell_borders(cell, color="000000", size=4):
    tcPr = cell._tc.get_or_add_tcPr()
    _remove(tcPr, "w:tcBorders")
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def _cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    _remove(tcPr, "w:tcMar")
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def _cell_valign(cell, align=WD_ALIGN_VERTICAL.CENTER):
    tcPr = cell._tc.get_or_add_tcPr()
    _remove(tcPr, "w:vAlign")
    va = OxmlElement("w:vAlign")
    va.set(qn("w:val"), "center" if align == WD_ALIGN_VERTICAL.CENTER else "top")
    tcPr.append(va)


def _col_width(cell, twips):
    tcPr = cell._tc.get_or_add_tcPr()
    _remove(tcPr, "w:tcW")
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _row_height(row, twips, exact=False):
    trPr = row._tr.get_or_add_trPr()
    _remove(trPr, "w:trHeight")
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"), str(int(twips)))
    trH.set(qn("w:hRule"), "exact" if exact else "atLeast")
    trPr.append(trH)


def _no_row_break(row):
    trPr = row._tr.get_or_add_trPr()
    _remove(trPr, "w:cantSplit")
    trPr.append(OxmlElement("w:cantSplit"))


def _style_cell(cell, shade=WHITE, border_color=BLACK, border_size=4, margins=True):
    _cell_shade(cell, shade)
    _cell_borders(cell, color=border_color, size=border_size)
    if margins:
        _cell_margins(cell)
    return cell


def _new_table(doc, rows, cols, col_widths, total_width):
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"

    tblPr = tbl._tbl.tblPr
    _remove(tblPr, "w:tblW")
    _remove(tblPr, "w:tblLayout")

    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(total_width))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)

    tblGrid = tbl._tbl.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        tbl._tbl.insert(0, tblGrid)

    for old in tblGrid.findall(qn("w:gridCol")):
        tblGrid.remove(old)

    for w in col_widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tblGrid.append(gc)

    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            _col_width(cell, col_widths[min(i, len(col_widths) - 1)])

    return tbl


# ──────────────────────────────────────────────────────────────────────────────
# PARAGRAPH WRITERS
# ──────────────────────────────────────────────────────────────────────────────

def _clear_to_single_para(cell):
    cell.text = ""
    para = cell.paragraphs[0]
    _para_spacing(para, 0, 0)
    return para


def _label_para(cell, text, size_pt=LABEL_PT, align=WD_ALIGN_PARAGRAPH.RIGHT):
    para = _clear_to_single_para(cell)
    para.alignment = align
    _para_spacing(para, 20, 20)
    for i, line in enumerate(text.split("\n")):
        if i:
            para.add_run("\n")
        styled_run(para, line, bold=True, size_pt=size_pt)
    return para


def _value_para(
    cell,
    text,
    size_pt=BASE_PT,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    bold=False,
    multiline=True,
):
    para = _clear_to_single_para(cell)
    para.alignment = align
    _para_spacing(para, 20, 20)

    value = text or ""
    if multiline and "\n" in value:
        for i, line in enumerate(value.split("\n")):
            if i:
                para.add_run("\n")
            styled_run(para, line, size_pt=size_pt, bold=bold)
    else:
        styled_run(para, value, size_pt=size_pt, bold=bold)
    return para


def _center_bold_para(cell, text, size_pt=BASE_PT, color_hex=BLACK):
    para = _clear_to_single_para(cell)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(para, 20, 20)
    styled_run(para, text, bold=True, size_pt=size_pt, color_hex=color_hex)
    return para


def _section_hdr_para(cell, text, size_pt=BODY_PT):
    para = _clear_to_single_para(cell)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _para_spacing(para, 20, 20)
    _keep_with_next(para, True)
    styled_run(para, text, bold=True, size_pt=size_pt)
    return para


def _content_para(cell, text="", before=30, after=20, size_pt=BODY_PT, bold_prefix=None, italic_prefix=None):
    para = cell.add_paragraph()
    _para_spacing(para, before, after)
    _keep_lines(para, True)

    if italic_prefix:
        styled_run(para, italic_prefix[0], bold=True, italic=True, size_pt=size_pt)
        styled_run(para, italic_prefix[1], size_pt=size_pt)
    elif bold_prefix:
        styled_run(para, bold_prefix[0], bold=True, size_pt=size_pt)
        styled_run(para, bold_prefix[1], size_pt=size_pt)
    else:
        styled_run(para, text, size_pt=size_pt)
    return para


def _heading_para(cell, text, size_pt=BODY_PT):
    para = cell.add_paragraph()
    _para_spacing(para, 40, 20)
    _keep_with_next(para, True)
    styled_run(para, text, bold=True, underline=True, size_pt=size_pt)
    return para


def _empty_para(cell, before=0, after=0):
    para = cell.add_paragraph()
    _para_spacing(para, before, after)
    return para


def _bullet_para(cell, text, is_sub=False):
    para = cell.add_paragraph()
    _para_spacing(para, 20, 20)
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(900 if is_sub else 540))
    ind.set(qn("w:hanging"), "260")
    pPr.append(ind)
    bullet_char = "\u25E6" if is_sub else "\u2022"
    styled_run(para, f"{bullet_char}  {text}", size_pt=BODY_PT)
    return para


def _semi_breaks(cell, text, size_pt=BODY_PT):
    segments = [s.strip() for s in str(text).split(";") if s.strip()]
    if not segments:
        _content_para(cell, "")
        return

    for i, seg in enumerate(segments):
        suffix = ";" if i < len(segments) - 1 else ""
        para = cell.add_paragraph()
        _para_spacing(para, 20, 20)
        styled_run(para, seg + suffix, size_pt=size_pt)


def _checkbox(checked: bool) -> str:
    return "\u2612" if checked else "\u2610"


# ──────────────────────────────────────────────────────────────────────────────
# NORMALIZERS
# ──────────────────────────────────────────────────────────────────────────────

def _norm_revision(entry):
    if isinstance(entry, dict):
        return (
            str(entry.get("date", "")),
            str(entry.get("version", "")),
            str(entry.get("updated_by", "")),
            str(entry.get("description", "")),
        )
    if isinstance(entry, (list, tuple)):
        padded = list(entry) + ["", "", "", ""]
        return tuple(str(x) for x in padded[:4])
    return (str(entry), "", "", "")


# ──────────────────────────────────────────────────────────────────────────────
# DOCUMENT SETUP
# ──────────────────────────────────────────────────────────────────────────────

def _setup_document(doc: Document):
    sec = doc.sections[0]
    sec.page_width = Inches(_PAGE_W_IN)
    sec.page_height = Inches(_PAGE_H_IN)
    sec.left_margin = Inches(_MARGIN_LR_IN)
    sec.right_margin = Inches(_MARGIN_LR_IN)
    sec.top_margin = Inches(_MARGIN_TOP_IN)
    sec.bottom_margin = Inches(_MARGIN_BOTTOM_IN)
    sec.footer_distance = Inches(0.3)

    doc.styles["Normal"].font.name = FONT_FAMILY
    doc.styles["Normal"].font.size = Pt(BASE_PT)
    doc.styles["Normal"].paragraph_format.space_before = Pt(0)
    doc.styles["Normal"].paragraph_format.space_after = Pt(0)


def _gap(doc, before=18, after=0):
    p = doc.add_paragraph()
    _para_spacing(p, before, after)


# ──────────────────────────────────────────────────────────────────────────────
# ZONE 1 — HEADER
# ──────────────────────────────────────────────────────────────────────────────

def render_header_banner(doc: Document, logo_path: str | None):
    table = _new_table(doc, 1, 1, [_CONTENT_W], _CONTENT_W)

    banner = table.rows[0].cells[0]
    _style_cell(banner, GRAY_BANNER)
    _cell_valign(banner, WD_ALIGN_VERTICAL.CENTER)

    # This controls the BOX size.
    _row_height(table.rows[0], 1280, exact=True)
    _no_row_break(table.rows[0])

    # Padding inside the banner so the logo sits centered with breathing room.
    _cell_margins(banner, top=100, bottom=100, left=0, right=0)

    banner.text = ""
    p = banner.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p, 0, 0)

    logo_ok = bool(logo_path and os.path.exists(str(logo_path)))
    if logo_ok:
        try:
            run = p.add_run()
            run.add_picture(str(logo_path), width=Inches(BANNER_LOGO_WIDTH_IN))
            return
        except Exception:
            logo_ok = False

    if not logo_ok:
        styled_run(p, "wipro", bold=True, color_hex=WIPRO_RED, size_pt=22)
        styled_run(p, ":", bold=True, color_hex=WIPRO_RED, size_pt=22)
        styled_run(p, "  healthplan services", color_hex=WIPRO_NAVY, size_pt=20)


# ──────────────────────────────────────────────────────────────────────────────
# ZONE 2 — METADATA GRID
# ──────────────────────────────────────────────────────────────────────────────

def render_metadata_zone(doc: Document, data: dict[str, Any]):
    W = _CONTENT_W
    CL = int(W * 0.18)
    CM = int(W * 0.33)
    CRL = int(W * 0.18)
    CRV = W - CL - CM - CRL

    top = _new_table(doc, 11, 4, [CL, CM, CRL, CRV], W)

    def _prep_label(cell, text):
        _style_cell(cell, GRAY_LABEL)
        _cell_valign(cell, WD_ALIGN_VERTICAL.CENTER)
        _cell_margins(cell, top=35, bottom=35, left=60, right=60)
        _label_para(cell, text)

    def _prep_value(cell, text, align=WD_ALIGN_PARAGRAPH.LEFT):
        _style_cell(cell, WHITE)
        _cell_valign(cell, WD_ALIGN_VERTICAL.CENTER)
        _cell_margins(cell, top=35, bottom=35, left=60, right=60)
        _value_para(cell, text, align=align)

    row_heights = {
        0: 300,   # Policy Name - slightly tighter so it tucks under the larger banner
        1: 340,
        2: 340,
        3: 420,
        4: 420,
        5: 620,   # Custodian row expanded
        6: 300,
        7: 320,
        8: 320,
        9: 700,
        10: 320,
    }

    for idx, twips in row_heights.items():
        _row_height(top.rows[idx], twips, exact=False)
        _no_row_break(top.rows[idx])

    _prep_label(top.cell(0, 0), "Policy Name")
    merged = top.cell(0, 1).merge(top.cell(0, 3))
    _prep_value(merged, _safe(data.get("policy_name")))

    _prep_label(top.cell(1, 0), "Policy Number")
    _prep_value(top.cell(1, 1), _safe(data.get("policy_number")))
    _prep_label(top.cell(1, 2), "Version Number")
    _prep_value(top.cell(1, 3), _safe(data.get("version")))

    _prep_label(top.cell(2, 0), "")
    _prep_value(top.cell(2, 1), "")
    _prep_label(top.cell(2, 2), "GRC ID Number")
    _prep_value(top.cell(2, 3), _safe(data.get("grc_id")))

    _prep_label(top.cell(3, 0), "Supersedes Policy")
    _prep_value(top.cell(3, 1), _safe(data.get("supersedes")))
    _prep_label(top.cell(3, 2), "Effective Date")
    _prep_value(top.cell(3, 3), _safe(data.get("effective_date")))

    _prep_label(top.cell(4, 0), "Last Reviewed\nDate")
    _prep_value(top.cell(4, 1), _safe(data.get("last_reviewed")))
    _prep_label(top.cell(4, 2), "Last Revised Date")
    _prep_value(top.cell(4, 3), _safe(data.get("last_revised")))

    _prep_label(top.cell(5, 0), "Policy Custodian\nName(s)")
    merged = top.cell(5, 1).merge(top.cell(5, 3))
    _style_cell(merged, WHITE)
    _cell_valign(merged, WD_ALIGN_VERTICAL.CENTER)
    _cell_margins(merged, top=65, bottom=65, left=60, right=60)
    _value_para(
        merged,
        _safe(data.get("custodians")),
        align=WD_ALIGN_PARAGRAPH.LEFT,
        multiline=True,
    )

    left_hdr = top.cell(6, 0).merge(top.cell(6, 1))
    right_hdr = top.cell(6, 2).merge(top.cell(6, 3))
    _style_cell(left_hdr, GRAY_SUBHDR)
    _style_cell(right_hdr, GRAY_SUBHDR)
    _cell_valign(left_hdr, WD_ALIGN_VERTICAL.CENTER)
    _cell_valign(right_hdr, WD_ALIGN_VERTICAL.CENTER)
    _center_bold_para(left_hdr, "Policy Owner")
    _center_bold_para(right_hdr, "Policy Approver")

    _prep_label(top.cell(7, 0), "Name")
    _prep_value(top.cell(7, 1), _safe(data.get("owner_name")))
    _prep_label(top.cell(7, 2), "Name")
    _prep_value(top.cell(7, 3), _safe(data.get("approver_name")))

    _prep_label(top.cell(8, 0), "Title")
    _prep_value(top.cell(8, 1), _safe(data.get("owner_title")))
    _prep_label(top.cell(8, 2), "Title")
    _prep_value(top.cell(8, 3), _safe(data.get("approver_title")))

    _prep_label(top.cell(9, 0), "Signature")
    _prep_value(top.cell(9, 1), "")
    _prep_label(top.cell(9, 2), "Signature")
    _prep_value(top.cell(9, 3), "")

    _prep_label(top.cell(10, 0), "Date Signed")
    _prep_value(top.cell(10, 1), _safe(data.get("date_signed")))
    _prep_label(top.cell(10, 2), "Date Approved")
    _prep_value(top.cell(10, 3), _safe(data.get("date_approved")))


# ──────────────────────────────────────────────────────────────────────────────
# ZONE 3 — APPLICABILITY / LOB MATRIX
# ──────────────────────────────────────────────────────────────────────────────

def render_applicability_zone(doc: Document, data: dict[str, Any]):
    W = _CONTENT_W
    app_to = _dict(data.get("applicable_to"))
    lob = _dict(data.get("line_of_business"))

    rows_def = [
        ("HealthPlan Services, Inc.", _bool(app_to.get("hps_inc"))),
        ("HealthPlan Services Insurance Agency, LLC", _bool(app_to.get("agency"))),
        ("Policy Types", None),
        ("Corporate", _bool(app_to.get("corporate"))),
        ("Government Affairs Review Required", _bool(app_to.get("govt_affairs"))),
        ("Legal Review Required", _bool(app_to.get("legal_review"))),
        ("Line of Business (LOB)", None),
        ("All LOBs", _bool(lob.get("all_lobs"))),
        (f"Specific LOB [{_safe(lob.get('specific_lob')) or 'INSERT HERE'}]", _bool(lob.get("specific_lob_checked"))),
    ]

    LEFT = int(W * 0.42)
    TEXT = int(W * 0.50)
    CHECK = W - LEFT - TEXT

    t2 = _new_table(doc, len(rows_def), 3, [LEFT, TEXT, CHECK], W)

    left_anchor = t2.rows[0].cells[0]
    for i in range(1, len(rows_def)):
        left_anchor = left_anchor.merge(t2.rows[i].cells[0])

    _style_cell(left_anchor, GRAY_LABEL)
    _cell_margins(left_anchor, top=160, bottom=160, left=120, right=120)
    _cell_valign(left_anchor, WD_ALIGN_VERTICAL.CENTER)
    left_anchor.text = ""

    lp = left_anchor.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(lp, 0, 0)
    styled_run(lp, "Applicable To:\n(select all that apply)", bold=True, size_pt=BODY_PT)

    row_heights = [360, 360, 220, 360, 360, 360, 220, 360, 360]

    for i, (label, checked) in enumerate(rows_def):
        row = t2.rows[i]
        _row_height(row, row_heights[i], exact=False)
        _no_row_break(row)

        text_cell = row.cells[1]
        check_cell = row.cells[2]
        is_header = checked is None

        _style_cell(text_cell, GRAY_SUBHDR if is_header else WHITE)
        _cell_margins(text_cell, top=40, bottom=40, left=70, right=70)
        _cell_valign(text_cell, WD_ALIGN_VERTICAL.CENTER)
        text_cell.text = ""

        tp = text_cell.paragraphs[0]
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.RIGHT
        _para_spacing(tp, 10, 10)
        if is_header:
            styled_run(tp, label, bold=True, size_pt=BASE_PT)
        else:
            styled_run(tp, label, size_pt=BASE_PT)

        _style_cell(check_cell, WHITE)
        _cell_margins(check_cell, top=20, bottom=20, left=10, right=10)
        _cell_valign(check_cell, WD_ALIGN_VERTICAL.CENTER)
        check_cell.text = ""

        cp = check_cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _para_spacing(cp, 0, 0)
        if not is_header:
            styled_run(cp, _checkbox(checked), size_pt=BASE_PT)


# ──────────────────────────────────────────────────────────────────────────────
# ZONE 4 — GENERIC SECTION BOX
# ──────────────────────────────────────────────────────────────────────────────

def render_section_box(doc: Document, heading: str, fill_fn):
    tbl = _new_table(doc, 2, 1, [_CONTENT_W], _CONTENT_W)

    hdr = tbl.rows[0].cells[0]
    _style_cell(hdr, GRAY_SECTION)
    _cell_margins(hdr, top=45, bottom=45, left=80, right=80)
    _section_hdr_para(hdr, heading)
    _no_row_break(tbl.rows[0])

    cnt = tbl.rows[1].cells[0]
    _style_cell(cnt, WHITE)
    _cell_margins(cnt, top=60, bottom=80, left=120, right=120)
    cnt.text = ""
    fill_fn(cnt)
    return tbl


# ──────────────────────────────────────────────────────────────────────────────
# CONTENT ZONES
# ──────────────────────────────────────────────────────────────────────────────

def render_purpose_zone(doc: Document, data: dict[str, Any]):
    def fill(cell):
        text = _safe(data.get("purpose"))
        if not text:
            _content_para(cell, "")
            return
        for line in text.split("\n"):
            line = line.strip()
            if line:
                _content_para(cell, line)
            else:
                _empty_para(cell)

    render_section_box(doc, "Purpose and Scope", fill)


def render_definitions_zone(doc: Document, data: dict[str, Any]):
    def fill(cell):
        defs = _dict(data.get("definitions"))
        if not defs:
            _content_para(cell, "")
            return
        for term, definition in defs.items():
            para = cell.add_paragraph()
            _para_spacing(para, 20, 20)
            styled_run(para, "\u2013  ", size_pt=BODY_PT)
            styled_run(para, f"{term}:  ", bold=True, size_pt=BODY_PT)
            styled_run(para, str(definition), size_pt=BODY_PT)

    render_section_box(doc, "Definitions", fill)


def render_policy_statement_zone(doc: Document, data: dict[str, Any]):
    def fill(cell):
        stmt = _safe(data.get("policy_statement"))
        para = cell.add_paragraph()
        _para_spacing(para, 30, 30)
        lower = stmt.lower()
        idx = lower.find(" that ") + len(" that ") if " that " in lower else 0
        if idx:
            styled_run(para, stmt[:idx], bold=True, italic=True, size_pt=BODY_PT)
            styled_run(para, stmt[idx:], size_pt=BODY_PT)
        else:
            styled_run(para, stmt, size_pt=BODY_PT)

    render_section_box(doc, "Policy Statement", fill)


def render_procedures_zone(doc: Document, data: dict[str, Any]):
    def fill(cell):
        procedures = _list(data.get("procedures"))
        if not procedures:
            _content_para(cell, "")
            return

        for item in procedures:
            if not isinstance(item, dict):
                _content_para(cell, _safe(item))
                continue

            kind = item.get("type", "para")
            text = _safe(item.get("text"))

            if kind == "empty":
                _empty_para(cell)
            elif kind == "heading":
                _heading_para(cell, text)
            elif kind == "bullet":
                _bullet_para(cell, text, is_sub=False)
            elif kind == "sub-bullet":
                _bullet_para(cell, text, is_sub=True)
            elif kind == "bold_intro":
                _content_para(
                    cell,
                    "",
                    bold_prefix=(_safe(item.get("bold")), _safe(item.get("rest"))),
                )
            elif kind == "bold_intro_semi":
                para = cell.add_paragraph()
                _para_spacing(para, 30, 20)
                styled_run(para, _safe(item.get("bold")), bold=True, size_pt=BODY_PT)
                segments = [s.strip() for s in _safe(item.get("rest")).split(";") if s.strip()]
                for i, seg in enumerate(segments):
                    suffix = ";" if i < len(segments) - 1 else ""
                    if i == 0:
                        styled_run(para, seg + suffix, size_pt=BODY_PT)
                    else:
                        p2 = cell.add_paragraph()
                        _para_spacing(p2, 0, 20)
                        styled_run(p2, seg + suffix, size_pt=BODY_PT)
            else:
                if ";" in text:
                    _semi_breaks(cell, text)
                else:
                    _content_para(cell, text)

    render_section_box(doc, "Procedures", fill)


def render_related_policies_zone(doc: Document, data: dict[str, Any]):
    def fill(cell):
        items = _list(data.get("related_policies"))
        if not items:
            _content_para(cell, "")
            return
        for item in items:
            _content_para(cell, _safe(item))

    render_section_box(doc, "Related Policies or Standard Operating Procedures", fill)


def render_citations_zone(doc: Document, data: dict[str, Any]):
    def fill(cell):
        items = _list(data.get("citations"))
        if not items:
            _content_para(cell, "")
            return
        for item in items:
            item_text = _safe(item)
            if ";" in item_text:
                _semi_breaks(cell, item_text)
            else:
                _content_para(cell, item_text)

    render_section_box(doc, "Citations/References", fill)


# ──────────────────────────────────────────────────────────────────────────────
# REVISION HISTORY
# ──────────────────────────────────────────────────────────────────────────────

def render_revision_history_zone(doc: Document, data: dict[str, Any]):
    rev_entries = []
    for entry in _list(data.get("revision_history")):
        try:
            rev_entries.append(_norm_revision(entry))
        except Exception:
            pass

    W = _CONTENT_W
    RC1 = int(W * 0.12)
    RC2 = int(W * 0.13)
    RC3 = int(W * 0.20)
    RC4 = W - RC1 - RC2 - RC3

    total_rows = 2 + max(1, len(rev_entries))
    rt = _new_table(doc, total_rows, 4, [RC1, RC2, RC3, RC4], W)

    rh = rt.rows[0].cells[0]
    rh = rh.merge(rt.rows[0].cells[3])
    _style_cell(rh, GRAY_SECTION)
    _cell_margins(rh, top=45, bottom=45, left=80, right=80)
    _section_hdr_para(rh, "Revision History")
    _no_row_break(rt.rows[0])

    headers = [
        ("Date", RC1),
        ("Version Number", RC2),
        ("Updated By", RC3),
        ("Description of Update", RC4),
    ]
    for ci, (hdr_text, w) in enumerate(headers):
        hc = rt.rows[1].cells[ci]
        _style_cell(hc, GRAY_LABEL)
        _col_width(hc, w)
        _value_para(hc, hdr_text, size_pt=LABEL_PT, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    _no_row_break(rt.rows[1])

    if not rev_entries:
        row = rt.rows[2]
        for ci, w in enumerate([RC1, RC2, RC3, RC4]):
            c = row.cells[ci]
            _style_cell(c, WHITE)
            _col_width(c, w)
            _value_para(c, "")
        return

    for ri, entry in enumerate(rev_entries, start=2):
        row = rt.rows[ri]
        for ci, (txt, w) in enumerate(zip(entry, [RC1, RC2, RC3, RC4])):
            rc = row.cells[ci]
            _style_cell(rc, WHITE)
            _col_width(rc, w)
            rc.text = ""
            lines = str(txt).split("\n")
            for li, line in enumerate(lines):
                para = rc.paragraphs[0] if li == 0 else rc.add_paragraph()
                _para_spacing(para, 15, 15)
                styled_run(para, line.strip(), size_pt=LABEL_PT)


# ──────────────────────────────────────────────────────────────────────────────
# FOOTER
# ──────────────────────────────────────────────────────────────────────────────

def render_footer(section, data: dict[str, Any]):
    footer = section.footer
    footer_tbl = footer.add_table(rows=2, cols=1, width=Inches(7.0))
    footer_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r in footer_tbl.rows:
        _no_row_break(r)
        c = r.cells[0]
        c.text = ""
        _cell_borders(c, color=WHITE, size=0)
        _cell_shade(c, WHITE)
        _cell_margins(c, top=0, bottom=0, left=0, right=0)
        _cell_valign(c, WD_ALIGN_VERTICAL.CENTER)

    fp1 = footer_tbl.rows[0].cells[0].paragraphs[0]
    fp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(fp1, 0, 0)
    styled_run(
        fp1,
        "Confidential & Proprietary \u00A9 HealthPlan Services Inc., including its subsidiaries and affiliates",
        size_pt=SMALL_PT,
        color_hex=FOOTER_GRAY,
    )

    fp2 = footer_tbl.rows[1].cells[0].paragraphs[0]
    fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(fp2, 0, 0)
    styled_run(
        fp2,
        f"{_safe(data.get('policy_number'))}  |  {_safe(data.get('policy_name'))}  |  {_safe(data.get('version'))}",
        size_pt=SMALL_PT,
        color_hex=FOOTER_GRAY,
    )


# ──────────────────────────────────────────────────────────────────────────────
# MAIN BUILDER
# ──────────────────────────────────────────────────────────────────────────────

def build_policy_document(data: dict, output_path: str, logo_path: str | None = None):
    doc = Document()
    _setup_document(doc)

    render_header_banner(doc, logo_path)
    render_metadata_zone(doc, data)
    _gap(doc)

    render_applicability_zone(doc, data)
    _gap(doc)

    render_purpose_zone(doc, data)
    _gap(doc)

    render_definitions_zone(doc, data)
    _gap(doc)

    render_policy_statement_zone(doc, data)
    _gap(doc)

    render_procedures_zone(doc, data)
    _gap(doc)

    render_related_policies_zone(doc, data)
    _gap(doc)

    render_citations_zone(doc, data)
    _gap(doc)

    render_revision_history_zone(doc, data)

    render_footer(doc.sections[0], data)

    doc.save(output_path)
    return output_path
