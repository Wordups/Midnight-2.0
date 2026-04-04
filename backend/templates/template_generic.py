"""
================================================================================
template_generic.py — Generic Minimal Policy Template
================================================================================
Design:
  - Clean black and white — no brand colors
  - Simple top header bar with optional logo (left, center, or right)
  - Industry-neutral field labels
  - No "Applicable To" checkboxes (HPS-specific, removed)
  - No "Policy Types" or "Line of Business" checkboxes (HPS-specific, removed)
  - Standard fields: Name, Number, Version, Classification, Status,
    Owner, Approver, Department, Effective Date, Review Date
  - Section headers: left-aligned bold with a thin black top border
  - Clean footer: policy number | name | version | classification
  - Compatible with any industry: healthcare, finance, tech, legal

Usage:
    from templates.template_generic import build_document
    build_document(policy_data, output_path, logo_path=None, logo_position="left")
================================================================================
"""

from __future__ import annotations

import os
import tempfile
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Template identity ─────────────────────────────────────────────────────────
TEMPLATE_NAME    = "Generic"
TEMPLATE_VERSION = "1.0"

# ── Page geometry ─────────────────────────────────────────────────────────────
_PAGE_W_IN  = 8.5
_MARGIN_IN  = 0.9
_CONTENT_W  = int((_PAGE_W_IN - 2 * _MARGIN_IN) * 1440)   # 9 720 twips

# ── Colors (minimal — black, white, light gray only) ─────────────────────────
BLACK       = "000000"
WHITE       = "FFFFFF"
LIGHT_GRAY  = "F2F2F2"
MID_GRAY    = "D9D9D9"
DARK_GRAY   = "595959"
HEADER_BG   = "1A1A1A"   # near-black header bar
HEADER_TEXT = "FFFFFF"


# ══════════════════════════════════════════════════════════════════════════════
# LOW-LEVEL XML HELPERS  (self-contained — no external imports)
# ══════════════════════════════════════════════════════════════════════════════

def _rgb(hex_str: str):
    h = hex_str.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def _remove(parent, tag):
    for old in parent.findall(qn(tag)):
        parent.remove(old)


def _run(para, text, bold=False, italic=False, underline=False,
         color=BLACK, size=10.0, font="Arial"):
    run = para.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.underline = underline
    run.font.name = font
    run.font.size = Pt(size)
    r, g, b = _rgb(color)
    run.font.color.rgb = RGBColor(r, g, b)
    return run


def _spacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    _remove(pPr, "w:spacing")
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:before"), str(before))
    spc.set(qn("w:after"),  str(after))
    pPr.append(spc)


def _keep_next(para, on=True):
    pPr = para._p.get_or_add_pPr()
    _remove(pPr, "w:keepNext")
    if on: pPr.append(OxmlElement("w:keepNext"))


def _keep_lines(para, on=True):
    pPr = para._p.get_or_add_pPr()
    _remove(pPr, "w:keepLines")
    if on: pPr.append(OxmlElement("w:keepLines"))


def _shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    _remove(tcPr, "w:shd")
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill.upper())
    tcPr.append(shd)


def _borders(cell, color=BLACK, size=4, sides=("top","left","bottom","right")):
    tcPr = cell._tc.get_or_add_tcPr()
    _remove(tcPr, "w:tcBorders")
    borders = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        el = OxmlElement(f"w:{side}")
        if side in sides:
            el.set(qn("w:val"),   "single")
            el.set(qn("w:sz"),    str(size))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
        else:
            el.set(qn("w:val"), "none")
        borders.append(el)
    tcPr.append(borders)


def _no_borders(cell):
    _borders(cell, color=WHITE, size=0, sides=())
    tcPr = cell._tc.get_or_add_tcPr()
    _remove(tcPr, "w:tcBorders")
    borders = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    tcPr.append(borders)


def _margins(cell, top=60, bottom=60, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    _remove(tcPr, "w:tcMar")
    mar = OxmlElement("w:tcMar")
    for side, val in (("top",top),("left",left),("bottom",bottom),("right",right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def _valign(cell, align=WD_ALIGN_VERTICAL.CENTER):
    tcPr = cell._tc.get_or_add_tcPr()
    _remove(tcPr, "w:vAlign")
    va = OxmlElement("w:vAlign")
    va.set(qn("w:val"), "center" if align == WD_ALIGN_VERTICAL.CENTER else "top")
    tcPr.append(va)


def _col_w(cell, twips):
    tcPr = cell._tc.get_or_add_tcPr()
    _remove(tcPr, "w:tcW")
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _row_h(row, twips, exact=False):
    trPr = row._tr.get_or_add_trPr()
    _remove(trPr, "w:trHeight")
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"),   str(int(twips)))
    trH.set(qn("w:hRule"), "exact" if exact else "atLeast")
    trPr.append(trH)


def _no_split(row):
    trPr = row._tr.get_or_add_trPr()
    _remove(trPr, "w:cantSplit")
    trPr.append(OxmlElement("w:cantSplit"))


def _new_table(doc, rows, cols, col_widths, total_width):
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style     = "Table Grid"

    tblPr = tbl._tbl.tblPr
    _remove(tblPr, "w:tblW"); _remove(tblPr, "w:tblLayout")

    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    str(total_width))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)

    tblGrid = tbl._tbl.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        tbl._tbl.insert(0, tblGrid)
    for old in tblGrid.findall(qn("w:gridCol")):
        tblGrid.remove(old)
    for w in col_widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tblGrid.append(gc)

    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            _col_w(cell, col_widths[min(i, len(col_widths)-1)])

    return tbl


# ══════════════════════════════════════════════════════════════════════════════
# REVISION HISTORY NORMALIZER
# ══════════════════════════════════════════════════════════════════════════════

def _norm_rev(entry):
    if isinstance(entry, dict):
        return (
            str(entry.get("date","")),
            str(entry.get("version","")),
            str(entry.get("updated_by","")),
            str(entry.get("description","")),
        )
    if isinstance(entry, (list, tuple)):
        padded = list(entry) + ["","","",""]
        return tuple(str(x) for x in padded[:4])
    return (str(entry),"","","")


# ══════════════════════════════════════════════════════════════════════════════
# CONTENT RENDERERS  (procedures, bullets, etc.)
# ══════════════════════════════════════════════════════════════════════════════

def _render_procedures(cell, procedures: list):
    for item in procedures:
        kind = item.get("type", "para")
        text = item.get("text", "")

        if kind == "empty":
            p = cell.add_paragraph()
            _spacing(p, 0, 0)

        elif kind == "heading":
            p = cell.add_paragraph()
            _spacing(p, 60, 20)
            _keep_next(p, True)
            _run(p, text, bold=True, underline=True, size=10.5)

        elif kind == "bullet":
            p = cell.add_paragraph()
            _spacing(p, 30, 30)
            _keep_lines(p, True)
            pPr = p._p.get_or_add_pPr()
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"),    "540")
            ind.set(qn("w:hanging"), "260")
            pPr.append(ind)
            _run(p, f"\u2022  {text}", size=10)

        elif kind == "sub-bullet":
            p = cell.add_paragraph()
            _spacing(p, 20, 20)
            pPr = p._p.get_or_add_pPr()
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"),    "900")
            ind.set(qn("w:hanging"), "260")
            pPr.append(ind)
            _run(p, f"\u25E6  {text}", size=10)

        elif kind in ("bold_intro", "bold_intro_semi"):
            bold_text = item.get("bold", "")
            rest_text = item.get("rest", "")

            if kind == "bold_intro_semi":
                p = cell.add_paragraph()
                _spacing(p, 40, 20)
                _run(p, bold_text, bold=True, size=10)
                segs = [s.strip() for s in rest_text.split(";") if s.strip()]
                for i, seg in enumerate(segs):
                    suffix = ";" if i < len(segs)-1 else ""
                    if i == 0:
                        _run(p, seg + suffix, size=10)
                    else:
                        p2 = cell.add_paragraph()
                        _spacing(p2, 0, 20)
                        _run(p2, seg + suffix, size=10)
            else:
                p = cell.add_paragraph()
                _spacing(p, 40, 30)
                _keep_lines(p, True)
                _run(p, bold_text, bold=True, size=10)
                _run(p, rest_text, size=10)

        else:  # para
            if ";" in text:
                segs = [s.strip() for s in text.split(";") if s.strip()]
                for i, seg in enumerate(segs):
                    p = cell.add_paragraph()
                    _spacing(p, 20, 20)
                    _run(p, seg + (";" if i < len(segs)-1 else ""), size=10)
            else:
                p = cell.add_paragraph()
                _spacing(p, 40, 30)
                _keep_lines(p, True)
                _run(p, text, size=10)


# ══════════════════════════════════════════════════════════════════════════════
# MAIN BUILDER
# ══════════════════════════════════════════════════════════════════════════════

def build_document(
    data:          dict,
    output_path:   str,
    logo_path:     str | None = None,
    logo_position: str        = "left",   # "left" | "center" | "right"
) -> str:
    """
    Build a generic minimal policy document.

    Args:
        data:          POLICY_DATA dict from extraction pipeline
        output_path:   where to save the .docx
        logo_path:     optional path to logo image
        logo_position: logo placement in header — left, center, or right

    Returns:
        output_path
    """
    doc = Document()
    W   = _CONTENT_W

    # ── Page setup ────────────────────────────────────────────────────────────
    sec = doc.sections[0]
    sec.page_width      = Inches(8.5)
    sec.page_height     = Inches(11)
    sec.left_margin     = Inches(_MARGIN_IN)
    sec.right_margin    = Inches(_MARGIN_IN)
    sec.top_margin      = Inches(0.75)
    sec.bottom_margin   = Inches(0.9)
    sec.footer_distance = Inches(0.3)

    doc.styles["Normal"].paragraph_format.space_before = Pt(0)
    doc.styles["Normal"].paragraph_format.space_after  = Pt(0)

    def _gap(before=36, after=0):
        p = doc.add_paragraph()
        _spacing(p, before, after)

    # ── Convenience: field values ─────────────────────────────────────────────
    policy_name    = data.get("policy_name",    "")
    policy_number  = data.get("policy_number",  "")
    version        = data.get("version",        "")
    grc_id         = data.get("grc_id",         "")
    supersedes     = data.get("supersedes",     "")
    effective_date = data.get("effective_date", "")
    last_reviewed  = data.get("last_reviewed",  "")
    last_revised   = data.get("last_revised",   "")
    custodians     = data.get("custodians",     "")
    owner_name     = data.get("owner_name",     "")
    owner_title    = data.get("owner_title",    "")
    approver_name  = data.get("approver_name",  "")
    approver_title = data.get("approver_title", "")
    date_signed    = data.get("date_signed",    "")
    date_approved  = data.get("date_approved",  "")

    # ──────────────────────────────────────────────────────────────────────────
    # HEADER BAR  — dark background, logo + org name
    # ──────────────────────────────────────────────────────────────────────────
    hdr_tbl = _new_table(doc, 1, 3, [W//3, W//3, W - 2*(W//3)], W)
    _row_h(hdr_tbl.rows[0], 720, exact=True)
    _no_split(hdr_tbl.rows[0])

    for cell in hdr_tbl.rows[0].cells:
        _shade(cell, HEADER_BG)
        _no_borders(cell)
        _margins(cell, top=0, bottom=0, left=120, right=120)
        _valign(cell, WD_ALIGN_VERTICAL.CENTER)
        cell.text = ""

    logo_ok = logo_path and os.path.exists(str(logo_path))
    logo_col_idx  = {"left": 0, "center": 1, "right": 2}.get(logo_position, 0)
    title_col_idx = 1 if logo_col_idx != 1 else 2

    # Logo cell
    logo_cell = hdr_tbl.rows[0].cells[logo_col_idx]
    logo_para = logo_cell.paragraphs[0]
    logo_para.alignment = {0: WD_ALIGN_PARAGRAPH.LEFT,
                           1: WD_ALIGN_PARAGRAPH.CENTER,
                           2: WD_ALIGN_PARAGRAPH.RIGHT}[logo_col_idx]
    _spacing(logo_para, 0, 0)

    if logo_ok:
        try:
            run = logo_para.add_run()
            run.add_picture(str(logo_path), width=Inches(1.8))
        except Exception:
            logo_ok = False

    if not logo_ok:
        _run(logo_para, "POLICY", bold=True, color=HEADER_TEXT, size=14)

    # Title cell
    title_cell = hdr_tbl.rows[0].cells[title_col_idx]
    title_para = title_cell.paragraphs[0]
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _spacing(title_para, 0, 0)
    _run(title_para, "POLICY DOCUMENT", bold=True, color=HEADER_TEXT, size=11)

    _gap(before=24)

    # ──────────────────────────────────────────────────────────────────────────
    # METADATA TABLE  — clean two-column layout
    # ──────────────────────────────────────────────────────────────────────────
    CL  = int(W * 0.22)
    CV  = int(W * 0.28)
    CL2 = int(W * 0.22)
    CV2 = W - CL - CV - CL2

    meta_rows = [
        ("Policy Title",      policy_name,    "",               ""),
        ("Policy Number",     policy_number,  "Version",        version),
        ("Classification",    "Internal Use", "Status",         "Active"),
        ("Effective Date",    effective_date, "Last Reviewed",  last_reviewed),
        ("Last Revised",      last_revised,   "Supersedes",     supersedes),
        ("GRC ID",            grc_id,         "Custodians",     custodians),
    ]

    meta_tbl = _new_table(doc, len(meta_rows), 4, [CL, CV, CL2, CV2], W)

    for ri, (l1, v1, l2, v2) in enumerate(meta_rows):
        row = meta_tbl.rows[ri]
        _no_split(row)

        cells = row.cells
        merge_row = l2 == ""

        # Label cells
        for ci in (0, 2):
            c = cells[ci]
            _shade(c, LIGHT_GRAY)
            _borders(c, color=MID_GRAY, size=4)
            _margins(c, top=60, bottom=60, left=80, right=80)
            _valign(c, WD_ALIGN_VERTICAL.CENTER)
            c.text = ""
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _spacing(p, 20, 20)
            _run(p, l1 if ci == 0 else l2, bold=True, size=9.0, color=DARK_GRAY)

        # Value cells
        if merge_row and ri == 0:
            # Policy title spans full width — merge value cells
            merged = cells[1].merge(cells[3])
            _shade(merged, WHITE)
            _borders(merged, color=MID_GRAY, size=4)
            _margins(merged, top=60, bottom=60, left=100, right=80)
            _valign(merged, WD_ALIGN_VERTICAL.CENTER)
            _col_w(merged, CV + CL2 + CV2)
            merged.text = ""
            p = merged.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _spacing(p, 20, 20)
            _run(p, v1 or "", size=9.5, bold=True)
        else:
            for ci in (1, 3):
                c = cells[ci]
                _shade(c, WHITE)
                _borders(c, color=MID_GRAY, size=4)
                _margins(c, top=60, bottom=60, left=100, right=80)
                _valign(c, WD_ALIGN_VERTICAL.CENTER)
                c.text = ""
                p = c.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _spacing(p, 20, 20)
                val = v1 if ci == 1 else v2
                _run(p, val or "", size=9.5)

    _gap()

    # ──────────────────────────────────────────────────────────────────────────
    # OWNERSHIP TABLE  — Owner | Approver
    # ──────────────────────────────────────────────────────────────────────────
    HL = int(W * 0.5)
    HR = W - HL

    own_tbl = _new_table(doc, 3, 2, [HL, HR], W)

    # Header row
    for ci, label in enumerate(["Policy Owner", "Policy Approver"]):
        c = own_tbl.rows[0].cells[ci]
        _shade(c, HEADER_BG)
        _borders(c, color=BLACK, size=4)
        _margins(c, top=50, bottom=50, left=100, right=80)
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(p, 0, 0)
        _run(p, label, bold=True, color=HEADER_TEXT, size=9.5)
    _no_split(own_tbl.rows[0])

    # Name row
    names = [f"{owner_name}", f"{approver_name}"]
    for ci, name in enumerate(names):
        c = own_tbl.rows[1].cells[ci]
        _shade(c, WHITE)
        _borders(c, color=MID_GRAY, size=4)
        _margins(c, top=60, bottom=60, left=100, right=80)
        c.text = ""
        p = c.paragraphs[0]
        _spacing(p, 20, 10)
        _run(p, name, bold=True, size=10)
    _no_split(own_tbl.rows[1])

    # Title + signature row
    titles = [
        f"{owner_title}\n\nSignature: ___________________________   Date: {date_signed or '___________'}",
        f"{approver_title}\n\nSignature: ___________________________   Date: {date_approved or '___________'}",
    ]
    for ci, txt in enumerate(titles):
        c = own_tbl.rows[2].cells[ci]
        _shade(c, WHITE)
        _borders(c, color=MID_GRAY, size=4)
        _margins(c, top=60, bottom=60, left=100, right=80)
        _row_h(own_tbl.rows[2], 720)
        c.text = ""
        lines = txt.split("\n")
        for li, line in enumerate(lines):
            p = c.paragraphs[0] if li == 0 else c.add_paragraph()
            _spacing(p, 20, 20)
            _run(p, line, size=9.0, color=DARK_GRAY if "Signature" in line else BLACK)
    _no_split(own_tbl.rows[2])

    _gap()

    # ──────────────────────────────────────────────────────────────────────────
    # SECTION HELPER  — clean top-border heading + content cell
    # ──────────────────────────────────────────────────────────────────────────
    def _section(heading: str, fill_fn):
        # Section heading paragraph (not a table — simpler, cleaner)
        h_para = doc.add_paragraph()
        _spacing(h_para, 24, 8)
        _keep_next(h_para, True)

        # Top border on heading paragraph
        pPr = h_para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        top_bdr = OxmlElement("w:top")
        top_bdr.set(qn("w:val"),   "single")
        top_bdr.set(qn("w:sz"),    "6")
        top_bdr.set(qn("w:space"), "4")
        top_bdr.set(qn("w:color"), BLACK)
        pBdr.append(top_bdr)
        pPr.append(pBdr)

        _run(h_para, heading.upper(), bold=True, size=10.5, color=BLACK)

        # Content table — single cell, white background
        tbl = _new_table(doc, 1, 1, [W], W)
        cnt = tbl.rows[0].cells[0]
        _shade(cnt, WHITE)
        _no_borders(cnt)
        _margins(cnt, top=40, bottom=60, left=0, right=0)
        cnt.text = ""
        fill_fn(cnt)

    # ── Purpose ───────────────────────────────────────────────────────────────
    def _purpose(cell):
        purpose = data.get("purpose", "") or ""
        for line in purpose.strip().split("\n"):
            line = line.strip()
            if line:
                p = cell.add_paragraph()
                _spacing(p, 30, 30)
                _keep_lines(p, True)
                _run(p, line, size=10)
            else:
                p = cell.add_paragraph()
                _spacing(p, 0, 0)

    _section("Purpose and Scope", _purpose)
    _gap()

    # ── Definitions ───────────────────────────────────────────────────────────
    def _definitions(cell):
        defs = data.get("definitions") or {}
        if not defs:
            p = cell.add_paragraph()
            _spacing(p, 30, 30)
            _run(p, "No definitions provided.", size=10, color=DARK_GRAY)
            return
        for term, defn in defs.items():
            p = cell.add_paragraph()
            _spacing(p, 30, 30)
            _run(p, f"{term}:  ", bold=True, size=10)
            _run(p, str(defn), size=10)

    _section("Definitions", _definitions)
    _gap()

    # ── Policy Statement ──────────────────────────────────────────────────────
    def _statement(cell):
        stmt = data.get("policy_statement", "") or ""
        if not stmt:
            return
        p = cell.add_paragraph()
        _spacing(p, 30, 30)
        lower = stmt.lower()
        idx = lower.find(" that ") + len(" that ") if " that " in lower else 0
        if idx:
            _run(p, stmt[:idx], bold=True, italic=True, size=10)
            _run(p, stmt[idx:], size=10)
        else:
            _run(p, stmt, size=10)

    _section("Policy Statement", _statement)
    _gap()

    # ── Procedures ────────────────────────────────────────────────────────────
    def _procedures(cell):
        procs = data.get("procedures") or []
        if not procs:
            p = cell.add_paragraph()
            _spacing(p, 30, 30)
            _run(p, "No procedures defined.", size=10, color=DARK_GRAY)
            return
        _render_procedures(cell, procs)

    _section("Procedures", _procedures)
    _gap()

    # ── Related Policies ──────────────────────────────────────────────────────
    def _related(cell):
        items = data.get("related_policies") or []
        if not items:
            p = cell.add_paragraph()
            _spacing(p, 30, 30)
            _run(p, "None.", size=10, color=DARK_GRAY)
            return
        for item in items:
            p = cell.add_paragraph()
            _spacing(p, 25, 25)
            _run(p, f"\u2022  {item}", size=10)

    _section("Related Policies", _related)
    _gap()

    # ── Citations ─────────────────────────────────────────────────────────────
    def _citations(cell):
        items = data.get("citations") or []
        if not items:
            p = cell.add_paragraph()
            _spacing(p, 30, 30)
            _run(p, "None.", size=10, color=DARK_GRAY)
            return
        for cit in items:
            if ";" in str(cit):
                segs = [s.strip() for s in str(cit).split(";") if s.strip()]
                for i, seg in enumerate(segs):
                    p = cell.add_paragraph()
                    _spacing(p, 20, 20)
                    _run(p, seg + (";" if i < len(segs)-1 else ""), size=10)
            else:
                p = cell.add_paragraph()
                _spacing(p, 25, 25)
                _run(p, str(cit), size=10)

    _section("Citations / References", _citations)
    _gap()

    # ──────────────────────────────────────────────────────────────────────────
    # REVISION HISTORY TABLE
    # ──────────────────────────────────────────────────────────────────────────
    rev_raw     = data.get("revision_history") or []
    rev_entries = []
    for entry in rev_raw:
        try:
            rev_entries.append(_norm_rev(entry))
        except Exception:
            pass

    # Section heading
    h_para = doc.add_paragraph()
    _spacing(h_para, 24, 8)
    _keep_next(h_para, True)
    pPr = h_para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top_bdr = OxmlElement("w:top")
    top_bdr.set(qn("w:val"),   "single")
    top_bdr.set(qn("w:sz"),    "6")
    top_bdr.set(qn("w:space"), "4")
    top_bdr.set(qn("w:color"), BLACK)
    pBdr.append(top_bdr)
    pPr.append(pBdr)
    _run(h_para, "REVISION HISTORY", bold=True, size=10.5)

    RC1 = int(W * 0.12)
    RC2 = int(W * 0.12)
    RC3 = int(W * 0.22)
    RC4 = W - RC1 - RC2 - RC3

    rev_tbl = _new_table(doc, 1 + max(len(rev_entries), 1), 4, [RC1, RC2, RC3, RC4], W)

    # Header row
    for ci, (hdr, w) in enumerate([("Date",RC1),("Version",RC2),("Updated By",RC3),("Description",RC4)]):
        c = rev_tbl.rows[0].cells[ci]
        _shade(c, LIGHT_GRAY)
        _borders(c, color=MID_GRAY, size=4)
        _margins(c, top=50, bottom=50, left=80, right=80)
        _col_w(c, w)
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(p, 0, 0)
        _run(p, hdr, bold=True, size=9.0)
    _no_split(rev_tbl.rows[0])

    if rev_entries:
        for ri, entry in enumerate(rev_entries, 1):
            row = rev_tbl.rows[ri]
            for ci, (txt, w) in enumerate(zip(entry, [RC1, RC2, RC3, RC4])):
                c = row.cells[ci]
                _shade(c, WHITE)
                _borders(c, color=MID_GRAY, size=4)
                _margins(c, top=50, bottom=50, left=80, right=80)
                _col_w(c, w)
                c.text = ""
                lines = str(txt).split("\n")
                for li, line in enumerate(lines):
                    p = c.paragraphs[0] if li == 0 else c.add_paragraph()
                    _spacing(p, 15, 15)
                    _run(p, line.strip(), size=9.0)
    else:
        row = rev_tbl.rows[1]
        c = row.cells[0].merge(row.cells[3])
        _shade(c, WHITE)
        _no_borders(c)
        _margins(c, top=50, bottom=50, left=80, right=80)
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(p, 0, 0)
        _run(p, "No revision history recorded.", size=9.5, color=DARK_GRAY)

    # ──────────────────────────────────────────────────────────────────────────
    # FOOTER
    # ──────────────────────────────────────────────────────────────────────────
    footer     = sec.footer
    footer_tbl = footer.add_table(rows=2, cols=1, width=Inches(6.7))
    footer_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri, (line, align) in enumerate([
        (f"{policy_number}  |  {policy_name}  |  {version}  |  Internal Use",
         WD_ALIGN_PARAGRAPH.CENTER),
        ("Confidential — For internal use only. Unauthorized distribution is prohibited.",
         WD_ALIGN_PARAGRAPH.CENTER),
    ]):
        c = footer_tbl.rows[ri].cells[0]
        c.text = ""
        _no_borders(c)
        _shade(c, WHITE)
        _margins(c, top=0, bottom=0, left=0, right=0)
        p = c.paragraphs[0]
        p.alignment = align
        _spacing(p, 0, 0)
        _run(p, line, size=7.5, color=DARK_GRAY)

    doc.save(output_path)
    return output_path
